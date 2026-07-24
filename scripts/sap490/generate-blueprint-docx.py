"""Fill the official SAP490 Blueprint template with IDTS EN/VI content.

The template under ``docs/sap490/templates`` is read-only.  Each output starts
as a byte-for-byte copy and retains its cover, sections, headers/footers,
styles, numbering, and eight core tables before approved content tables are added.
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
import shutil

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "sap490" / "generated"
TEMPLATE = ROOT / "docs" / "sap490" / "templates" / "Deliverable_template" / "Blueprint_Template.docx"
VERSION = "v0.4"
DATE = "2026-07-23"


EN = {
    "language": "English",
    "output": "Blueprint_IDTS_SAP01_en_v0.4.docx",
    "title": "Issue and Defect Tracking System in SAP",
    "subtitle": "SAP490 Business Blueprint",
    "status": "Draft for mentor review",
    "prepared": "DonHV - Project Leader / BA-PM Consolidation",
    "document_control": "Document Control",
    "version_history": "Version History",
    "approval": "Review and Approval",
    "toc": "Table of Contents",
    "toc_cache": "Live table of contents - refresh fields in Microsoft Word before final submission.",
    "toc_note": "The live TOC field is configured to refresh when the document opens in Microsoft Word.",
    "page_label": "Page",
    "overview": "1. Blueprint Overview",
    "purpose": "1.1 Purpose and business context",
    "purpose_text": (
        "IDTS is an internal SAP-oriented issue and defect tracking system for software testing. "
        "It provides a controlled process to record, classify, assign, review, follow up, retest, "
        "close, audit, and monitor defects without becoming a source-code or full ALM platform."
    ),
    "objectives": "1.2 Business objectives",
    "objectives_items": [
        "Reduce duplicate or incomplete defect reports through structured capture and similar-bug support.",
        "Assign defects using Application Component, Defect Category, Component Category, and Developer Responsibility.",
        "Keep assignee and current action owner clear throughout rejection and information-request follow-up.",
        "Preserve comments, history, notifications, and reviewable evidence for mentor and QA review.",
        "Give PM users workload, overdue, queue, status, and nextProcessor visibility.",
    ],
    "baseline": "1.3 Current solution baseline",
    "scope": "2. Scope and Organization",
    "in_scope": "2.1 In scope",
    "out_scope": "2.2 Out of scope",
    "roles": "2.3 Roles and responsibilities",
    "process": "3. Business Process Blueprint",
    "process_intro": (
        "The normal create flow persists Assigned when an authorized user explicitly selects a valid Developer, "
        "or Pending Assignment when no suitable Developer is selected. New remains legacy/import compatibility only."
    ),
    "process_flow": "3.1 End-to-end flow",
    "status_rules": "3.2 Lifecycle and ownership rules",
    "capabilities": "4. Functional Capability Blueprint",
    "data": "5. Information and Integration Blueprint",
    "entities": "5.1 Core information objects",
    "integrations": "5.2 Runtime profiles and integrations",
    "reports": "6. Reports and Monitoring",
    "quality": "7. Security, Quality, and Control",
    "traceability": "8. Traceability and Evidence",
    "limitations": "9. Known Limitations and Pending Acceptance",
    "glossary": "10. Glossary",
    "footer": "IDTS SAP01 | SAP490 Blueprint | English",
    "version_summary": (
        "Promoted to v0.2 as an IDTS-native, source-generated Blueprint; removed inherited sample company/author/process content; "
        "aligned current CAP/Fiori, authentication, PostgreSQL/S3, email-outbox, and advisory-AI baselines."
    ),
    "approval_rows": [
        ["Prepared by", "DonHV", "Project Leader / BA-PM consolidation", "Drafted", DATE],
        ["Reviewed by", "Mentor / Supervisor", "Review scope, evidence, and SAP490 fit", "Pending", "TBD"],
        ["Approved by", "Mentor / Supervisor", "Approve Blueprint baseline", "Pending", "TBD"],
    ],
    "baseline_rows": [
        ["Application", "SAP CAP Node.js, OData V4, SAP Fiori Elements/SAPUI5"],
        ["Authentication", "AuthService login with server-managed AuthSessions; backend role checks remain authoritative"],
        ["Data profiles", "SQLite for local development; Render-hosted shared QA uses PostgreSQL through the integration profile"],
        ["Attachments", "@cap-js/attachments with local DB fallback and externally bound S3 object storage for shared QA"],
        ["Notifications", "In-app Notifications plus asynchronous NotificationDeliveries email outbox; provider failure does not roll back workflow"],
        ["Advisory AI", "Optional suggestions stored as normalized AiSuggestions audit rows; human review is mandatory and normal CAP actions remain authoritative"],
    ],
    "in_scope_items": [
        "Structured defect creation, duplicate support, SAP-context classification, assignment and reassignment.",
        "Pending Assignment, Developer review, request more information, reject with reason and clear follow-up ownership.",
        "Status processing through Resolved, Retest Required, Closed, and Reopened.",
        "Comments, attachments, history events/logs, in-app notifications, email outbox, and PM monitoring.",
        "Optional review-only AI suggestions for similar bugs, classification, summaries, and Smart Assign rationale.",
    ],
    "out_scope_items": [
        "Direct source-code fixing, source control, code review, CI/CD, transport, or release management.",
        "A full Jira, SAP Cloud ALM, SAP Solution Manager, ServiceNow, or enterprise incident-management replacement.",
        "Autonomous AI assignment, classification, duplicate confirmation, lifecycle transitions, or mandatory root-cause analysis.",
        "Hardcoded credentials, private endpoints, organization-wide integrations, or a final production topology decision.",
    ],
    "role_rows": [
        ["Tester", "Create and classify bugs; select assignee when appropriate; provide requested information; retest; close or reopen; comment and attach evidence."],
        ["Developer", "Review assigned/authorized bugs; comment; request information; reject unsuitable assignment/classification with reason; progress and resolve."],
        ["PM", "Monitor all defects, workload, overdue items, pending/rejected/retest queues, nextProcessor, history, reports, and escalation notifications."],
        ["System", "Validate authorization and business rules; maintain status and nextProcessor; persist audit, notification, delivery, attachment, and AI-review records."],
    ],
    "process_rows": [
        ["1", "Detect and check", "Tester detects a defect and searches for an existing open/similar bug.", "Tester"],
        ["2", "Create and classify", "Capture reproduction, expected/actual result, priority, severity, environment, Application Component, Defect Category, and optional SAP Module.", "Tester"],
        ["3", "Assign or queue", "Select a valid Developer explicitly, or submit to Pending Assignment. The system does not auto-pick a Developer during create.", "Tester / PM / System"],
        ["4", "Review and process", "Developer moves through In Review/In Progress, requests information, rejects unsuitable work, or resolves with required context.", "Developer"],
        ["5", "Follow up", "Tester/PM resubmits requested information or corrects/reassigns Rejected work. Rejected is not a terminal state.", "Tester / PM"],
        ["6", "Retest and close", "Tester/PM verifies resolved work through Retest Required, then closes or reopens.", "Tester / PM"],
        ["7", "Audit and monitor", "System records history/notifications; PM monitors workload, overdue items, queues, and ownership.", "System / PM"],
    ],
    "status_rows": [
        ["New", "Legacy/import compatibility only", "Not persisted by the normal create flow"],
        ["Pending Assignment", "Tester/PM or unassigned queue", "No suitable Developer explicitly selected"],
        ["Assigned / In Review / In Progress", "Assigned Developer", "Technical handling states"],
        ["Need More Information", "Tester/PM follow-up", "Developer reason and nextProcessor required"],
        ["Rejected", "Tester/PM follow-up", "Reason and correction/reassignment path required; not terminal"],
        ["Resolved / Retest Required", "Tester/PM verification", "Retest before final acceptance when required"],
        ["Closed / Reopened", "Tester/PM", "Close accepted work; reopen when the issue persists"],
    ],
    "capability_rows": [
        ["BP-01", "Bug reporting", "Create structured defects with human-readable bug number and evidence context.", "Must"],
        ["BP-02", "Duplicate support", "Search and optionally review similar-bug suggestions; confirmed relationships remain a human decision.", "Must"],
        ["BP-03", "Classification", "Maintain optional SAP Module and required Application Component/Defect Category with valid Component Category.", "Must"],
        ["BP-04", "Assignment", "Filter by Developer Responsibility; validate selected Developer; support Pending Assignment and reassignment.", "Must"],
        ["BP-05", "Lifecycle", "Enforce role/status transitions, required reasons, nextProcessor, retest, close, and reopen.", "Must"],
        ["BP-06", "Collaboration", "Comments and attachments support clarification and QA evidence without directly changing status.", "Must"],
        ["BP-07", "Audit and notification", "Persist grouped history, field-level logs, in-app notifications, and separate email-delivery attempts.", "Must"],
        ["BP-08", "PM monitoring", "Expose workload, overdue, pending, rejected, retest, and current-action-owner views.", "Must"],
        ["BP-09", "Advisory AI", "Provide optional review-only suggestions with safe audit data and failure isolation.", "Optional"],
    ],
    "entity_rows": [
        ["Bugs", "Defect record, classification, status, ownership, dates, reproduction, expected/actual result"],
        ["Users / Developers / AuthSessions", "Internal profile, business role, Developer view, and server-managed authentication session"],
        ["SAPModules / ApplicationComponents / DefectCategories / ComponentCategories", "Classification master data and valid assignment key"],
        ["DeveloperResponsibilities", "Developer capability mapping by Component Category and optional SAP Module"],
        ["Comments / Attachments", "Discussion and evidence; attachment content follows configured storage profile"],
        ["HistoryEvents / HistoryLogs", "Readable grouped history and append-only field-level audit"],
        ["Notifications / NotificationDeliveries", "In-app event and separate asynchronous email-delivery outbox"],
        ["DuplicateLinks / AiSuggestions", "Confirmed human-reviewed relationships and normalized advisory-AI audit rows"],
    ],
    "integration_rows": [
        ["Local", "SQLite and local attachment DB fallback", "Developer verification only; not production acceptance"],
        ["Shared QA", "Render CAP runtime, PostgreSQL integration profile, externally bound S3 object storage", "Binding and persistence evidence required"],
        ["Email", "Database outbox processed asynchronously through private SMTP configuration", "No live-provider success claim without approved evidence"],
        ["AI", "Disabled-by-default mock or optional server-side provider", "Failure-safe; no secrets, attachments, raw response, or autonomous decision"],
    ],
    "report_rows": [
        ["Operational defect list", "Status, priority, severity, classification, assignee, nextProcessor, due/overdue", "Tester / Developer / PM according to role"],
        ["PM workload and queue view", "Workload, Pending Assignment, Rejected follow-up, Retest Required, overdue", "PM"],
        ["Audit evidence", "History events/logs, comments, attachments, notification and delivery status", "Authorized users / mentor evidence"],
        ["Test and defect evidence", "Test Scenario, Unit Test, Functional Test, Test Report, Test and Fix Bug", "QA / mentor review; evidence status must remain truthful"],
    ],
    "quality_rows": [
        ["Authorization", "Backend resolves bearer sessions and enforces Tester, Developer, and PM permissions; UI visibility is not the security boundary."],
        ["Data integrity", "Required classification, valid responsibility, required reasons, lifecycle transitions, and nextProcessor rules are server-validated."],
        ["Auditability", "Important changes retain actor, role, timestamp, action, old/new value, and reason where applicable."],
        ["Failure isolation", "Email or AI provider failure cannot roll back or block the normal authorized bug workflow."],
        ["No-secret handling", "Passwords, bearer tokens, SMTP credentials, object-storage secrets, private endpoints, and AI credentials remain outside source and review artifacts."],
        ["Usability", "Fiori List Report/Object Page, value helps, messages, semantic status, and role-aware actions support beginner and enterprise review."],
    ],
    "trace_rows": [
        ["Business requirements", "docs/ba/brd/brd.en.md", "BRD v1.5"],
        ["Software requirements", "docs/ba/srs/srs.en.md", "SRS v1.4"],
        ["Functional behavior", "docs/ba/frs/frs.en.md", "FRS v1.5"],
        ["Business rules and scope", "IDTS-Business-Rule.md; IDTS-PROJECT-SCOPE-SAP01.md", "Canonical"],
        ["Architecture and process", "docs/diagrams and rendered Diagram Pack", "Source-controlled Mermaid/PlantUML plus rendered assets"],
        ["QA evidence", "Current mentor Test Scenario, Unit Test, Functional Test, Test Report, Test and Fix Bug", "Only executed evidence may be reported as PASS"],
    ],
    "limitation_items": [
        "Mentor review and approval remain Pending; this Draft is not a signed acceptance artifact.",
        "UAT is Prepared but not signed. The Final Project Report remains a template/pending deliverable and is not claimed complete.",
        "Live SMTP/provider success is not claimed without approved private configuration and captured evidence.",
        "Final shared-QA attachment acceptance depends on active PostgreSQL/S3 bindings and persistence/download evidence.",
        "Optional AI is disabled by default; live-provider capability is not a prerequisite for the normal workflow and is not claimed without evidence.",
        "The current educational/shared-QA baseline is not a final enterprise production topology, availability commitment, or performance certification.",
    ],
    "glossary_rows": [
        ["Assignee", "Developer who owns technical handling of the bug"],
        ["nextProcessor", "Current person or role/queue expected to take the next action; not a second assignee"],
        ["Component Category", "Valid Application Component and Defect Category pair used for assignment"],
        ["Developer Responsibility", "Mapping of Developer capability to Component Category and optional SAP Module"],
        ["NotificationDeliveries", "Separate asynchronous email-delivery outbox and attempt record"],
        ["AiSuggestions", "Normalized, safe, human-review audit record for optional advisory AI output"],
    ],
}


VI = {
    "language": "Vietnamese",
    "output": "Blueprint_IDTS_SAP01_vi_v0.4.docx",
    "title": "Hệ thống Quản lý Issue và Defect trong SAP",
    "subtitle": "Business Blueprint SAP490",
    "status": "Bản nháp chờ người hướng dẫn đánh giá",
    "prepared": "DonHV — Trưởng dự án / Tổng hợp BA-PM",
    "document_control": "Kiểm soát tài liệu",
    "version_history": "Lịch sử phiên bản",
    "approval": "Đánh giá và phê duyệt",
    "toc": "Mục lục",
    "toc_cache": "Mục lục được cập nhật trong Microsoft Word trước khi phát hành bản cuối.",
    "toc_note": "Các trường mục lục được cấu hình để cập nhật khi mở tài liệu trong Microsoft Word.",
    "page_label": "Trang",
    "overview": "1. Tổng quan Blueprint",
    "purpose": "1.1 Mục đích và bối cảnh nghiệp vụ",
    "purpose_text": (
        "IDTS là hệ thống nội bộ theo định hướng SAP dùng để quản lý issue và defect trong hoạt động kiểm thử phần mềm. "
        "Hệ thống hỗ trợ quy trình có kiểm soát để ghi nhận, phân loại, phân công, đánh giá, xử lý tiếp, kiểm thử lại, "
        "đóng, lưu vết và giám sát defect; IDTS không phải nền tảng quản lý mã nguồn hoặc hệ thống ALM đầy đủ."
    ),
    "objectives": "1.2 Mục tiêu nghiệp vụ",
    "objectives_items": [
        "Giảm báo cáo defect trùng lặp hoặc thiếu thông tin bằng biểu mẫu có cấu trúc và chức năng hỗ trợ tìm bug tương tự.",
        "Phân công defect theo Application Component, Defect Category, Component Category và Developer Responsibility.",
        "Làm rõ người phụ trách kỹ thuật và người phải thực hiện bước tiếp theo trong các luồng từ chối hoặc yêu cầu bổ sung thông tin.",
        "Duy trì bình luận, lịch sử, thông báo và bằng chứng có thể kiểm tra cho người hướng dẫn và QA.",
        "Cung cấp cho PM thông tin về khối lượng công việc, quá hạn, hàng đợi, trạng thái và nextProcessor.",
    ],
    "baseline": "1.3 Nền giải pháp hiện tại",
    "scope": "2. Phạm vi và tổ chức",
    "in_scope": "2.1 Trong phạm vi",
    "out_scope": "2.2 Ngoài phạm vi",
    "roles": "2.3 Vai trò và trách nhiệm",
    "process": "3. Blueprint quy trình nghiệp vụ",
    "process_intro": (
        "Luồng tạo thông thường lưu trạng thái Assigned khi người dùng có quyền chủ động chọn Developer hợp lệ; "
        "nếu chưa chọn được Developer phù hợp, defect được lưu ở trạng thái Pending Assignment. "
        "New chỉ được giữ để tương thích với dữ liệu cũ hoặc dữ liệu nhập."
    ),
    "process_flow": "3.1 Luồng đầu-cuối",
    "status_rules": "3.2 Quy tắc vòng đời và quyền sở hữu",
    "capabilities": "4. Blueprint năng lực chức năng",
    "data": "5. Blueprint thông tin và tích hợp",
    "entities": "5.1 Các đối tượng thông tin cốt lõi",
    "integrations": "5.2 Môi trường chạy và tích hợp",
    "reports": "6. Báo cáo và giám sát",
    "quality": "7. An toàn, chất lượng và kiểm soát",
    "traceability": "8. Truy vết và bằng chứng",
    "limitations": "9. Giới hạn đã biết và nội dung chờ nghiệm thu",
    "glossary": "10. Thuật ngữ",
    "footer": "IDTS SAP01 | SAP490 Blueprint | Tiếng Việt",
    "version_summary": (
        "Phiên bản v0.3 được điền từ template Blueprint chính thức; nội dung IDTS được đồng bộ với CAP/Fiori, "
        "AuthService/AuthSessions, PostgreSQL/S3, email outbox và phạm vi AI tư vấn hiện có."
    ),
    "approval_rows": [
        ["Người lập", "DonHV", "Trưởng dự án / Tổng hợp BA-PM", "Đã chuẩn bị", DATE],
        ["Người đánh giá", "Mentor / Supervisor", "Đánh giá phạm vi, bằng chứng và mức phù hợp SAP490", "Đang chờ", ""],
        ["Người phê duyệt", "Mentor / Supervisor", "Phê duyệt nền Blueprint", "Đang chờ", ""],
    ],
    "baseline_rows": [
        ["Ứng dụng", "SAP CAP Node.js, OData V4, SAP Fiori Elements/SAPUI5"],
        ["Xác thực", "AuthService đăng nhập với AuthSessions do máy chủ quản lý; kiểm tra vai trò ở backend là lớp quyết định cuối"],
        ["Hồ sơ dữ liệu", "SQLite dùng cho phát triển cục bộ; môi trường QA dùng chung trên Render sử dụng PostgreSQL qua cấu hình integration"],
        ["Tệp đính kèm", "@cap-js/attachments dùng cơ sở dữ liệu cục bộ dự phòng và kho đối tượng S3 được liên kết ngoài cho môi trường QA dùng chung"],
        ["Thông báo", "Notifications trong ứng dụng và hộp thư đi bất đồng bộ NotificationDeliveries; lỗi nhà cung cấp không hoàn tác quy trình nghiệp vụ"],
        ["AI tư vấn", "Gợi ý tùy chọn được lưu thành bản ghi AiSuggestions đã chuẩn hóa ở trạng thái PENDING; CAP vẫn quyết định mọi thao tác nghiệp vụ"],
    ],
    "in_scope_items": [
        "Tạo defect có cấu trúc, hỗ trợ kiểm tra trùng lặp, phân loại theo ngữ cảnh SAP, phân công và phân công lại.",
        "Pending Assignment, Developer đánh giá, yêu cầu bổ sung thông tin, từ chối có lý do và người xử lý tiếp theo rõ ràng.",
        "Xử lý trạng thái qua Resolved, Retest Required, Closed và Reopened.",
        "Bình luận, tệp đính kèm, HistoryEvents/HistoryLogs, thông báo trong ứng dụng, hộp thư đi email và giám sát của PM.",
        "Gợi ý AI chỉ để tham khảo cho bug tương tự, phân loại, tóm tắt và giải thích Smart Assign.",
    ],
    "out_scope_items": [
        "Sửa mã nguồn trực tiếp, quản lý mã nguồn, quy trình duyệt mã, CI/CD, transport hoặc quản lý phát hành.",
        "Thay thế đầy đủ Jira, SAP Cloud ALM, SAP Solution Manager, ServiceNow hoặc hệ thống quản lý sự cố cấp doanh nghiệp.",
        "AI tự động phân công, phân loại, xác nhận trùng lặp, chuyển trạng thái hoặc bắt buộc phân tích nguyên nhân gốc.",
        "Ghi cứng thông tin xác thực, endpoint riêng, tích hợp toàn tổ chức hoặc quyết định kiến trúc vận hành chính thức cuối cùng.",
    ],
    "role_rows": [
        ["Tester", "Tạo và phân loại defect; chọn người phụ trách khi phù hợp; bổ sung thông tin; kiểm thử lại; đóng hoặc mở lại; bình luận và đính kèm bằng chứng."],
        ["Developer", "Đánh giá defect được phân công hoặc được phép xem; bình luận; yêu cầu thông tin; từ chối phân công/phân loại không phù hợp có lý do; xử lý và giải quyết."],
        ["PM", "Giám sát defect, khối lượng công việc, quá hạn, các hàng đợi Pending Assignment/Rejected/Retest Required, nextProcessor, lịch sử, báo cáo và thông báo nâng cấp."],
        ["Hệ thống", "Kiểm tra phân quyền và quy tắc nghiệp vụ; duy trì trạng thái/nextProcessor; lưu nhật ký, thông báo, giao nhận, tệp đính kèm và bản ghi AI tư vấn."],
    ],
    "process_rows": [
        ["1", "Phát hiện và kiểm tra", "Tester phát hiện defect và tìm bug đang mở hoặc bug tương tự đã tồn tại.", "Tester"],
        ["2", "Tạo và phân loại", "Nhập bước tái hiện, kết quả mong đợi/thực tế, mức ưu tiên, mức độ nghiêm trọng, môi trường, Application Component, Defect Category và SAP Module nếu cần.", "Tester"],
        ["3", "Phân công hoặc đưa vào hàng đợi", "Chủ động chọn Developer hợp lệ hoặc gửi vào Pending Assignment. Hệ thống không tự chọn Developer khi tạo defect.", "Tester / PM / Hệ thống"],
        ["4", "Đánh giá và xử lý", "Developer chuyển sang In Review/In Progress, yêu cầu thông tin, từ chối công việc không phù hợp hoặc chuyển Resolved với thông tin bắt buộc.", "Developer"],
        ["5", "Xử lý tiếp", "Tester/PM gửi lại thông tin hoặc sửa và phân công lại defect ở trạng thái Rejected. Rejected không phải trạng thái kết thúc.", "Tester / PM"],
        ["6", "Kiểm thử lại và đóng", "Tester/PM xác minh defect đã xử lý qua Retest Required, sau đó đóng hoặc mở lại.", "Tester / PM"],
        ["7", "Lưu vết và giám sát", "Hệ thống ghi lịch sử/thông báo; PM giám sát khối lượng, quá hạn, hàng đợi và quyền sở hữu.", "Hệ thống / PM"],
    ],
    "status_rows": [
        ["New", "Chỉ để tương thích dữ liệu cũ hoặc dữ liệu nhập", "Luồng tạo thông thường không lưu New"],
        ["Pending Assignment", "Tester/PM hoặc hàng đợi chưa phân công", "Chưa chủ động chọn được Developer phù hợp"],
        ["Assigned / In Review / In Progress", "Developer được phân công", "Các trạng thái xử lý kỹ thuật"],
        ["Need More Information", "Tester/PM xử lý tiếp", "Cần lý do của Developer và nextProcessor"],
        ["Rejected", "Tester/PM xử lý tiếp", "Cần lý do và đường sửa/phân công lại; không phải trạng thái kết thúc"],
        ["Resolved / Retest Required", "Tester/PM xác minh", "Kiểm thử lại trước khi nghiệm thu cuối nếu cần"],
        ["Closed / Reopened", "Tester/PM", "Đóng khi chấp nhận kết quả; mở lại khi lỗi vẫn tồn tại"],
    ],
    "capability_rows": [
        ["BP-01", "Báo cáo bug", "Tạo defect có cấu trúc, mã bug dễ đọc và ngữ cảnh bằng chứng.", "Bắt buộc"],
        ["BP-02", "Hỗ trợ kiểm tra trùng lặp", "Tìm và xem gợi ý bug tương tự; quan hệ xác nhận vẫn do con người quyết định.", "Bắt buộc"],
        ["BP-03", "Phân loại", "Duy trì SAP Module tùy chọn, Application Component/Defect Category bắt buộc và Component Category hợp lệ.", "Bắt buộc"],
        ["BP-04", "Phân công", "Lọc theo Developer Responsibility; kiểm tra Developer được chọn; hỗ trợ Pending Assignment và phân công lại.", "Bắt buộc"],
        ["BP-05", "Vòng đời", "Áp dụng quyền/trạng thái chuyển tiếp, lý do bắt buộc, nextProcessor, kiểm thử lại, đóng và mở lại.", "Bắt buộc"],
        ["BP-06", "Cộng tác", "Bình luận và tệp đính kèm hỗ trợ làm rõ và bằng chứng QA mà không trực tiếp đổi trạng thái.", "Bắt buộc"],
        ["BP-07", "Nhật ký và thông báo", "Lưu lịch sử nhóm, nhật ký cấp trường, thông báo trong ứng dụng và lần thử gửi email riêng.", "Bắt buộc"],
        ["BP-08", "Giám sát PM", "Cung cấp các góc nhìn về khối lượng, quá hạn, chờ phân công, từ chối, kiểm thử lại và người xử lý hiện tại.", "Bắt buộc"],
        ["BP-09", "AI tư vấn", "Cung cấp gợi ý tùy chọn chỉ để tham khảo, với dữ liệu nhật ký an toàn và cô lập lỗi.", "Tùy chọn"],
    ],
    "entity_rows": [
        ["Bugs", "Bản ghi defect, phân loại, trạng thái, quyền sở hữu, ngày, bước tái hiện và kết quả mong đợi/thực tế"],
        ["Users / Developers / AuthSessions", "Hồ sơ nội bộ, vai trò nghiệp vụ, góc nhìn Developer và phiên xác thực do máy chủ quản lý"],
        ["SAPModules / ApplicationComponents / DefectCategories / ComponentCategories", "Dữ liệu chủ phân loại và khóa phân công hợp lệ"],
        ["DeveloperResponsibilities", "Ánh xạ năng lực Developer theo Component Category và SAP Module tùy chọn"],
        ["Comments / Attachments", "Trao đổi và bằng chứng; nội dung tệp đính kèm tuân theo cấu hình lưu trữ"],
        ["HistoryEvents / HistoryLogs", "Lịch sử nhóm dễ đọc và nhật ký cấp trường chỉ được nối thêm"],
        ["Notifications / NotificationDeliveries", "Sự kiện trong ứng dụng và hộp thư đi email bất đồng bộ riêng"],
        ["DuplicateLinks / AiSuggestions", "Quan hệ trùng lặp đã được con người xác nhận và bản ghi AI tư vấn đã chuẩn hóa"],
    ],
    "integration_rows": [
        ["Cục bộ", "SQLite và cơ sở dữ liệu dự phòng cho tệp đính kèm", "Chỉ dùng để xác minh khi phát triển; không phải nghiệm thu môi trường vận hành chính thức"],
        ["QA dùng chung", "CAP trên Render, cấu hình PostgreSQL integration và kho đối tượng S3 được liên kết ngoài", "Cần bằng chứng liên kết và lưu bền"],
        ["Email", "Hộp thư đi trong cơ sở dữ liệu được xử lý bất đồng bộ qua cấu hình SMTP riêng", "Không tuyên bố gửi thành công qua nhà cung cấp thực khi chưa có bằng chứng được duyệt"],
        ["AI", "Mô phỏng mặc định tắt hoặc nhà cung cấp phía máy chủ tùy chọn", "Cô lập lỗi; không gửi bí mật, tệp đính kèm, phản hồi thô hoặc tự quyết định"],
    ],
    "report_rows": [
        ["Danh sách defect vận hành", "Trạng thái, ưu tiên, mức độ nghiêm trọng, phân loại, assignee, nextProcessor, đến hạn/quá hạn", "Tester / Developer / PM theo vai trò"],
        ["Góc nhìn khối lượng và hàng đợi của PM", "Khối lượng, Pending Assignment, xử lý tiếp Rejected, Retest Required và quá hạn", "PM"],
        ["Bằng chứng nhật ký", "HistoryEvents/HistoryLogs, bình luận, tệp đính kèm, trạng thái thông báo và giao nhận", "Người dùng được phép / người hướng dẫn"],
        ["Bằng chứng kiểm thử và defect", "Test Scenario, Unit Test, Functional Test, Test Report, Test and Fix Bug", "QA / người hướng dẫn; chỉ bằng chứng đã chạy mới được báo PASS"],
    ],
    "quality_rows": [
        ["Phân quyền", "Backend xác định phiên bearer và áp dụng quyền Tester, Developer, PM; khả năng nhìn thấy trên giao diện không phải ranh giới bảo mật."],
        ["Toàn vẹn dữ liệu", "Phân loại bắt buộc, trách nhiệm hợp lệ, lý do bắt buộc, chuyển trạng thái và nextProcessor được máy chủ kiểm tra."],
        ["Khả năng kiểm toán", "Thay đổi quan trọng giữ người thực hiện, vai trò, thời điểm, thao tác, giá trị cũ/mới và lý do khi áp dụng."],
        ["Cô lập lỗi", "Lỗi nhà cung cấp email hoặc AI không hoàn tác hay chặn quy trình bug hợp lệ."],
        ["Không lưu bí mật", "Mật khẩu, bearer token, thông tin SMTP, bí mật kho đối tượng, endpoint riêng và thông tin xác thực AI nằm ngoài mã nguồn/tài liệu đánh giá."],
        ["Khả năng sử dụng", "Fiori List Report/Object Page, value help, thông báo, trạng thái ngữ nghĩa và thao tác theo vai trò hỗ trợ việc đánh giá."],
    ],
    "trace_rows": [
        ["Yêu cầu nghiệp vụ", "docs/ba/brd/brd.vi.md", "BRD v1.5"],
        ["Yêu cầu phần mềm", "docs/ba/srs/srs.vi.md", "SRS v1.4"],
        ["Hành vi chức năng", "docs/ba/frs/frs.vi.md", "FRS v1.5"],
        ["Quy tắc và phạm vi nghiệp vụ", "IDTS-Business-Rule.md; IDTS-PROJECT-SCOPE-SAP01.md", "Nguồn chuẩn"],
        ["Kiến trúc và quy trình", "docs/diagrams và Diagram Pack đã kết xuất", "Nguồn Mermaid/PlantUML được quản lý trong repository cùng tài sản đã kết xuất"],
        ["Bằng chứng QA", "Bộ Test Scenario, Unit Test, Functional Test, Test Report và Test and Fix Bug hiện hành", "Chỉ bằng chứng đã chạy được báo PASS"],
    ],
    "limitation_items": [
        "Việc đánh giá và phê duyệt của người hướng dẫn vẫn đang chờ; bản nháp này không phải tài liệu nghiệm thu đã ký.",
        "UAT mới ở trạng thái Prepared và chưa ký xác nhận. Final Project Report vẫn là mẫu đang chờ, chưa được tuyên bố hoàn tất.",
        "Không tuyên bố gửi email thành công qua nhà cung cấp thực khi chưa có cấu hình riêng được duyệt và bằng chứng.",
        "Nghiệm thu cuối cho tệp đính kèm trên môi trường QA dùng chung phụ thuộc liên kết PostgreSQL/S3 đang hoạt động và bằng chứng lưu/tải.",
        "AI tùy chọn mặc định tắt; hệ thống hiện chưa có thao tác cập nhật trạng thái đánh giá AiSuggestions từ PENDING sang ACCEPTED, REJECTED hoặc IGNORED.",
        "Nền giáo dục/QA dùng chung hiện tại không phải kiến trúc vận hành doanh nghiệp cuối, cam kết sẵn sàng hoặc chứng nhận hiệu năng.",
    ],
    "glossary_rows": [
        ["Assignee", "Developer chịu trách nhiệm xử lý kỹ thuật chính của bug"],
        ["nextProcessor", "Người hoặc vai trò/hàng đợi phải thực hiện thao tác tiếp theo; không phải assignee thứ hai"],
        ["Component Category", "Cặp Application Component và Defect Category hợp lệ dùng cho phân công"],
        ["Developer Responsibility", "Ánh xạ năng lực Developer theo Component Category và SAP Module tùy chọn"],
        ["NotificationDeliveries", "Hộp thư đi email và bản ghi lần thử giao nhận bất đồng bộ riêng"],
        ["AiSuggestions", "Bản ghi AI tư vấn an toàn đã chuẩn hóa; trạng thái hiện tại được lưu là PENDING"],
    ],
}

# Detailed Blueprint content is kept as paired EN/VI structures so the two
# deliverables remain aligned by section, process ID, table shape, and meaning.
EN.update({
    "organization": "2.4 Organizational operating model",
    "organization_text": (
        "IDTS uses a small cross-functional operating model. Testers own defect quality and verification; "
        "Developers own authorized technical handling; PM owns queue health, escalation, and oversight; "
        "the System enforces permissions, state changes, audit, and delivery instructions. DonHV consolidates "
        "BA/PM and SAP490 evidence, while mentor review remains an external review activity rather than an approval already obtained."
    ),
    "raci": "2.5 RACI matrix",
    "raci_rows": [
        ["Create and classify defect", "R/A", "C", "C", "I"],
        ["Select or change assignee", "R", "C", "A", "I"],
        ["Review, request information, reject, resolve", "C", "R/A", "I", "I"],
        ["Resubmit requested information", "R", "C", "A", "I"],
        ["Retest, close, or reopen", "R", "C", "A", "I"],
        ["Maintain history and notifications", "I", "I", "C", "R/A"],
        ["Monitor workload, overdue, and queues", "C", "I", "R/A", "I"],
        ["Review AI suggestions and decide whether to continue with a normal authorized action", "R", "R", "A", "I"],
    ],
    "process_details": "3.3 Detailed business processes",
    "process_detail_intro": (
        "Each process below defines the business boundary that reviewers can trace to requirements, runtime behavior, and evidence. "
        "A status effect is stated only where the process is allowed to change lifecycle state."
    ),
    "data_ownership": "5.2 Data ownership and retention boundary",
    "interfaces": "5.3 Interface and integration catalogue",
    "embedded_diagrams": "5.4 Embedded process and architecture views",
    "diagram_note": "The figures are generated from the canonical sources under docs/diagrams and are included for review context.",
    "acceptance_trace": "8.1 Process-to-requirement-and-evidence traceability",
})

VI.update({
    "organization": "2.4 Mô hình vận hành tổ chức",
    "organization_text": (
        "IDTS vận hành theo mô hình liên chức năng gọn. Tester chịu trách nhiệm về chất lượng defect và việc xác minh; "
        "Developer chịu trách nhiệm xử lý kỹ thuật trong phạm vi được phép; PM giám sát hàng đợi, khối lượng công việc và các trường hợp cần xử lý tiếp. "
        "Hệ thống thực thi phân quyền, chuyển trạng thái, lưu vết và chỉ dẫn giao thông báo. DonHV tổng hợp BA/PM và bằng chứng SAP490; "
        "việc người hướng dẫn đánh giá vẫn đang chờ và không được xem là đã phê duyệt."
    ),
    "raci": "2.5 Ma trận RACI",
    "raci_rows": [
        ["Tạo và phân loại defect", "R/A", "C", "C", "I"],
        ["Chọn hoặc thay đổi người được phân công", "R", "C", "A", "I"],
        ["Đánh giá, yêu cầu thông tin, từ chối hoặc xác nhận đã xử lý", "C", "R/A", "I", "I"],
        ["Bổ sung và gửi lại thông tin được yêu cầu", "R", "C", "A", "I"],
        ["Kiểm thử lại, đóng hoặc mở lại", "R", "C", "A", "I"],
        ["Duy trì lịch sử và thông báo", "I", "I", "C", "R/A"],
        ["Giám sát khối lượng, quá hạn và hàng đợi", "C", "I", "R/A", "I"],
        ["Đánh giá gợi ý AI và quyết định có tiếp tục bằng thao tác nghiệp vụ được phép hay không", "R", "R", "A", "I"],
    ],
    "process_details": "3.3 Quy trình nghiệp vụ chi tiết",
    "process_detail_intro": (
        "Mỗi quy trình dưới đây xác định ranh giới nghiệp vụ để người đánh giá truy vết tới yêu cầu, hành vi thực tế của hệ thống và bằng chứng. "
        "Ảnh hưởng tới trạng thái chỉ được nêu khi quy trình thực sự được phép thay đổi trạng thái vòng đời."
    ),
    "data_ownership": "5.2 Quyền sở hữu dữ liệu và ranh giới lưu giữ",
    "interfaces": "5.3 Danh mục giao diện và tích hợp",
    "embedded_diagrams": "5.4 Sơ đồ quy trình và kiến trúc được nhúng",
    "diagram_note": "Các hình được tạo từ nguồn chuẩn trong docs/diagrams và được nhúng để phục vụ việc đánh giá.",
    "acceptance_trace": "8.1 Truy vết từ quy trình tới yêu cầu và bằng chứng",
})

EN_PROCESS_DETAILS = [
    ("BP-01", "Authentication and access", "Establish an authenticated IDTS identity and backend role boundary.", "Active user email and password supplied to AuthService.", "Bearer token backed by a server-side AuthSessions row.", "User; System", "Password hash verification, active-user check, session expiry/revocation, backend role mapping.", "No defect status change.", "Authenticated user or rejected login response.", "Authentication/session audit data only; never raw passwords or tokens.", "No defect notification."),
    ("BP-02", "Defect creation", "Create a complete, uniquely identifiable defect without inventing an assignee.", "Title, description, reproduction, actual/expected result, priority, severity, environment, classification and optional evidence.", "Persisted Bugs row with bugNumber, ownership, child evidence metadata, history and notification instructions.", "Tester; PM may coordinate", "Required fields, valid catalog codes, authorized creator, system-managed fields overwritten, selected assignee validated.", "Assigned when a valid Developer is explicitly selected; otherwise Pending Assignment. New is not the normal persisted start.", "Assigned Developer or Tester/PM queue.", "Create event plus field-level audit for authoritative values.", "In-app assignment/queue notification; email outbox instruction when eligible."),
    ("BP-03", "Duplicate and similar support", "Help a human avoid duplicate work without autonomous linking.", "Draft or persisted defect text and safe classification fields.", "Ranked candidates and optional suggested relation label.", "Tester or authorized reviewer", "Exclude source bug, bounded result count, deterministic fallback, sanitized AI/provider boundary.", "No automatic status change and no automatic DuplicateLinks write.", "Human reviews, ignores, or confirms a relationship separately.", "Safe AiSuggestions audit row only for a persisted source check.", "None unless a later confirmed business action requires it."),
    ("BP-04", "SAP-context classification", "Capture the business context used for reporting and responsibility matching.", "Optional SAP Module, Application Component, Defect Category.", "Validated Component Category and consistent classification references.", "Tester; PM correction support", "Application Component and Defect Category compatibility; SAP Module remains optional for pure IDTS defects.", "No status change by classification alone; may enable assignment/reassignment.", "Assignment decision owner.", "Old/new classification values and correction reason when material.", "Relevant assignee/follow-up notification after an actual assignment change."),
    ("BP-05", "Assignment and Pending Assignment", "Route work to an eligible Developer or an explicit follow-up queue.", "Validated classification, Developer Responsibilities, availability/workload context, optional selected Developer.", "Assignee and nextProcessor, or a clear unassigned queue.", "Tester/PM", "No automatic Developer selection during create; selected Developer must match active responsibility rules.", "Assigned when selected; Pending Assignment when no suitable Developer is selected; reassignment is an action, not a status.", "Assigned Developer or Tester/PM queue.", "Assignment/reassignment event with previous and new owner.", "Notify new assignee; PM/Tester visibility for pending queue."),
    ("BP-06", "Developer review and processing", "Allow the assigned Developer to assess and process authorized work.", "Assigned bug, evidence, comments, history and current capabilities.", "In Review/In Progress progression, request, rejection, or resolution decision.", "Assigned Developer", "Assigned-developer rule and transition allow-list; normal review notes optional, decision reasons required where specified.", "Assigned → In Review/In Progress; In Review may progress or resolve according to runtime.", "Assigned Developer until a follow-up decision changes owner.", "Status event and meaningful notes/reasons.", "Notify stakeholders when the action changes ownership or requires follow-up."),
    ("BP-07", "Request More Information and resubmit", "Obtain missing evidence, then return the same assigned work to Developer review.", "Developer reason; Tester/PM information update and required resubmit summary.", "Updated defect/comments/attachments and a restored Assigned state.", "Developer requests; Tester/PM responds", "Request action only from Assigned/In Review/In Progress by assigned Developer; resubmit requires existing assignee and update summary.", "Need More Information → Assigned on Resubmit to Developer. It is not a post-retest transition.", "Tester/PM while waiting; existing assigned Developer after resubmit.", "Request and resubmit events, reason, summary and changed fields.", "Notify Tester/PM on request and Developer on resubmit."),
    ("BP-08", "Rejection, correction and reassignment", "Return unsuitable classification or assignment for controlled correction rather than terminate the defect.", "Developer rejection reason and current assignment/classification.", "Rejected follow-up record, corrected data, reassignment or Pending Assignment.", "Developer rejects; Tester/PM corrects", "Reason mandatory; nextProcessor mandatory; Rejected cannot silently end and does not route through Need More Information in MVP.", "Assigned/In Review/In Progress → Rejected → Assigned or Pending Assignment.", "Tester/PM after rejection; new Developer or queue after correction.", "Immutable rejection reason plus correction/reassignment history.", "Notify Tester/PM, then the new assignee when work is reassigned."),
    ("BP-09", "Resolve, retest, close and reopen", "Separate Developer resolution from Tester/PM acceptance.", "Resolution note, current state, retest outcome or reopen reason.", "Resolved, Retest Required, Closed, or Reopened state with ownership.", "Developer resolves; Tester/PM verifies", "Resolution and reopen reasons required where configured; transitions must match the runtime state machine.", "Resolved → Retest Required/Closed/Reopened; Retest Required → Closed/Reopened. No Request More Information branch exists here.", "Tester/PM during verification; Developer after authorized reopen assignment flow.", "Resolution, retest, close or reopen event with actor and reason.", "Notify verification owner and assigned Developer when reopened."),
    ("BP-10", "Comments and attachments", "Preserve collaboration and evidence without bypassing lifecycle controls.", "Authorized comment text or attachment metadata/content through the configured storage profile.", "Comment or attachment linked to the defect and available to authorized readers.", "Tester, Developer, PM according to permission", "Comment role, file constraints, safe metadata, storage binding, and no-secret rule.", "No direct status change.", "Current lifecycle owner remains unchanged.", "Append collaboration/evidence history where applicable.", "Notify interested users only according to configured event rules."),
    ("BP-11", "Notifications and email outbox", "Record in-app events and attempt email delivery without coupling provider failure to workflow success.", "Committed business event and eligible active recipient.", "Notifications row plus NotificationDeliveries status/attempt data.", "System", "Recipient eligibility, private provider configuration, bounded retry, sanitized failure details.", "No independent status change; delivery failure never rolls back the business transaction.", "Recipient or operational retry owner.", "Delivery status, attempts, timing and safe error classification.", "In-app is persisted; email may be SENT, FAILED or SKIPPED. Local disabled provider remains SKIPPED."),
    ("BP-12", "PM monitoring and escalation", "Give PM an operational view of workload, queues, overdue items and current action ownership.", "Defect states, assignee, nextProcessor, due dates, workload and history.", "Role-aware workload, overdue, pending, rejected and retest views.", "PM", "Backend-filtered data and consistent current-action-owner calculation.", "Monitoring alone does not change status; authorized coordination actions use normal transitions.", "PM decides follow-up or escalation without becoming a second assignee.", "Any subsequent assignment/status action is audited normally.", "Escalation or queue notifications only when configured."),
    ("BP-13", "AI suggestions for human consideration", "Offer bounded assistance while keeping every business decision human and CAP-authoritative.", "Minimum allowlisted defect fields; no credentials, private email, attachments or storage references.", "Sanitized suggestion and confidence/explanation; a persisted audit row starts in PENDING.", "Authorized user", "AI disabled by default, timeout/failure isolation, output validation, no raw prompt/response persistence.", "No autonomous assignment, classification, duplicate confirmation or lifecycle change.", "The user reviews the displayed suggestion and may separately perform a normal authorized business action. The current service exposes no suggestion-state update action.", "AiSuggestions records the normalized request/result audit in PENDING; current actions do not persist ACCEPTED, REJECTED or IGNORED reviewer outcomes.", "No notification is created by the suggestion itself."),
]

VI_PROCESS_DETAILS = [
    ("BP-01", "Xác thực và truy cập", "Xác lập danh tính và ranh giới vai trò ở backend.", "Email và mật khẩu của người dùng đang hoạt động.", "Bearer token gắn với bản ghi AuthSessions phía máy chủ.", "Người dùng; Hệ thống", "Kiểm tra mật khẩu băm, trạng thái người dùng, hạn phiên và vai trò ở backend.", "Không đổi trạng thái defect.", "Người dùng đã xác thực hoặc phản hồi từ chối đăng nhập.", "Chỉ lưu vết phiên; không lưu mật khẩu hoặc token thô.", "Không phát sinh thông báo defect."),
    ("BP-02", "Tạo defect", "Tạo defect đầy đủ, có mã duy nhất và không tự suy diễn người được phân công.", "Tiêu đề, mô tả, bước tái hiện, kết quả thực tế/dự kiến, mức độ, môi trường, phân loại và bằng chứng tùy chọn.", "Bản ghi Bugs, mã bug, quyền sở hữu, lịch sử và chỉ dẫn thông báo.", "Tester; PM điều phối khi cần", "Kiểm tra trường bắt buộc, danh mục, quyền tạo và người được chọn.", "Assigned khi chọn Developer hợp lệ; nếu không là Pending Assignment. New không phải trạng thái khởi tạo thông thường.", "Developer được phân công hoặc hàng đợi Tester/PM.", "Sự kiện tạo và thay đổi các giá trị có thẩm quyền.", "Thông báo trong ứng dụng; tạo chỉ dẫn email khi đủ điều kiện."),
    ("BP-03", "Hỗ trợ phát hiện trùng lặp", "Hỗ trợ con người tránh xử lý trùng mà không tự tạo liên kết.", "Nội dung defect và các trường phân loại an toàn.", "Danh sách ứng viên đã xếp hạng và nhãn quan hệ gợi ý.", "Tester hoặc người đánh giá được phép", "Loại bug nguồn, giới hạn kết quả, phương án dự phòng xác định và dữ liệu nhà cung cấp đã làm sạch.", "Không tự đổi trạng thái hoặc ghi DuplicateLinks.", "Con người tự xác nhận quan hệ bằng thao tác riêng.", "Chỉ ghi AiSuggestions an toàn cho bug nguồn đã lưu.", "Không có, trừ khi thao tác nghiệp vụ tiếp theo yêu cầu."),
    ("BP-04", "Phân loại theo ngữ cảnh SAP", "Ghi ngữ cảnh phục vụ báo cáo và đối chiếu trách nhiệm.", "SAP Module tùy chọn, Application Component và Defect Category.", "Component Category hợp lệ và tham chiếu phân loại nhất quán.", "Tester; PM hỗ trợ hiệu chỉnh", "Kiểm tra tính tương thích giữa Application Component và Defect Category.", "Phân loại riêng không đổi trạng thái nhưng có thể cho phép phân công.", "Người chịu trách nhiệm quyết định phân công.", "Giá trị cũ/mới và lý do hiệu chỉnh khi cần.", "Chỉ thông báo sau thay đổi phân công thực tế."),
    ("BP-05", "Phân công và Pending Assignment", "Chuyển việc tới Developer phù hợp hoặc hàng đợi xử lý tiếp rõ ràng.", "Phân loại hợp lệ, Developer Responsibilities, tải công việc và Developer tùy chọn.", "Assignee và nextProcessor, hoặc hàng đợi chưa phân công.", "Tester/PM", "Không tự chọn Developer khi tạo; người được chọn phải có trách nhiệm đang hoạt động.", "Assigned khi chọn hợp lệ; Pending Assignment khi chưa có người phù hợp; phân công lại không phải trạng thái.", "Developer được phân công hoặc hàng đợi Tester/PM.", "Lưu sự kiện phân công với chủ sở hữu cũ và mới.", "Thông báo người mới; PM/Tester thấy hàng đợi chờ."),
    ("BP-06", "Developer đánh giá và xử lý", "Cho Developer được phân công xử lý trong phạm vi được phép.", "Defect đã phân công, bằng chứng, bình luận và lịch sử.", "Quyết định In Review/In Progress, yêu cầu, từ chối hoặc xác nhận đã xử lý.", "Developer được phân công", "Kiểm tra đúng người được phân công, chuyển trạng thái cho phép và lý do bắt buộc.", "Assigned → In Review/In Progress; tiếp tục theo máy trạng thái hiện hành.", "Developer giữ trách nhiệm tới khi quyết định tiếp theo đổi người xử lý.", "Lưu trạng thái và lý do có ý nghĩa.", "Thông báo khi thay đổi trách nhiệm hoặc cần xử lý tiếp."),
    ("BP-07", "Yêu cầu và gửi lại thông tin", "Thu thập bằng chứng thiếu rồi trả việc cho cùng Developer đánh giá.", "Lý do yêu cầu; thông tin bổ sung và tóm tắt gửi lại.", "Defect đã cập nhật và trở lại Assigned.", "Developer yêu cầu; Tester/PM phản hồi", "Chỉ Developer được phân công được yêu cầu; gửi lại cần assignee và tóm tắt.", "Need More Information → Assigned khi gửi lại.", "Tester/PM khi chờ; Developer hiện hữu sau khi gửi lại.", "Lưu yêu cầu, lý do, tóm tắt và thay đổi trường.", "Thông báo hai phía tại từng bước."),
    ("BP-08", "Từ chối, hiệu chỉnh và phân công lại", "Trả phân loại hoặc phân công sai về bước hiệu chỉnh có kiểm soát.", "Lý do từ chối và dữ liệu phân công/phân loại hiện tại.", "Thông tin đã sửa, phân công lại hoặc Pending Assignment.", "Developer từ chối; Tester/PM sửa", "Bắt buộc có lý do và nextProcessor; Rejected không được kết thúc im lặng.", "Assigned/In Review/In Progress → Rejected → Assigned hoặc Pending Assignment.", "Tester/PM sau từ chối; Developer mới hoặc hàng đợi sau hiệu chỉnh.", "Lưu lý do từ chối và lịch sử hiệu chỉnh.", "Thông báo Tester/PM rồi người được phân công mới."),
    ("BP-09", "Xác nhận xử lý, kiểm thử lại, đóng và mở lại", "Tách việc Developer xác nhận xử lý khỏi việc Tester/PM chấp nhận.", "Ghi chú xử lý, trạng thái, kết quả kiểm thử lại hoặc lý do mở lại.", "Resolved, Retest Required, Closed hoặc Reopened cùng quyền sở hữu.", "Developer xử lý; Tester/PM xác minh", "Lý do và chuyển trạng thái phải khớp máy trạng thái thực tế.", "Resolved → Retest Required/Closed/Reopened; Retest Required → Closed/Reopened.", "Tester/PM khi xác minh; Developer sau luồng mở lại được phép.", "Lưu sự kiện cùng người thực hiện và lý do.", "Thông báo người xác minh và Developer khi mở lại."),
    ("BP-10", "Bình luận và tệp đính kèm", "Lưu trao đổi và bằng chứng mà không bỏ qua kiểm soát vòng đời.", "Nội dung bình luận hoặc tệp qua cấu hình lưu trữ.", "Bình luận/tệp gắn với defect và chỉ người có quyền đọc được.", "Tester, Developer, PM theo quyền", "Kiểm tra vai trò, giới hạn tệp, siêu dữ liệu an toàn và cấu hình lưu trữ.", "Không trực tiếp đổi trạng thái.", "Người chịu trách nhiệm vòng đời không đổi.", "Bổ sung lịch sử trao đổi/bằng chứng khi áp dụng.", "Chỉ thông báo theo quy tắc sự kiện đã cấu hình."),
    ("BP-11", "Thông báo và hộp thư đi", "Ghi thông báo và thử gửi email mà không để lỗi nhà cung cấp làm hỏng quy trình.", "Sự kiện nghiệp vụ đã commit và người nhận hợp lệ.", "Notifications và dữ liệu trạng thái/lần thử trong NotificationDeliveries.", "Hệ thống", "Kiểm tra người nhận, cấu hình riêng tư, số lần thử và chi tiết lỗi đã làm sạch.", "Không tự đổi trạng thái; lỗi gửi không hoàn tác giao dịch nghiệp vụ.", "Người nhận hoặc người phụ trách thử lại.", "Lưu trạng thái giao, số lần thử, thời gian và loại lỗi an toàn.", "Trong ứng dụng được lưu; email là SENT, FAILED hoặc SKIPPED."),
    ("BP-12", "PM giám sát và điều phối", "Cho PM thấy khối lượng, hàng đợi, quá hạn và người xử lý hành động hiện tại.", "Trạng thái, assignee, nextProcessor, hạn, tải và lịch sử.", "Khung nhìn theo vai trò cho khối lượng, quá hạn và hàng đợi.", "PM", "Backend lọc dữ liệu và tính người xử lý hiện tại nhất quán.", "Giám sát không tự đổi trạng thái; điều phối dùng thao tác được phép.", "PM quyết định bước tiếp theo mà không trở thành assignee thứ hai.", "Các thay đổi phân công/trạng thái sau đó được lưu vết.", "Chỉ gửi cảnh báo khi đã cấu hình."),
    ("BP-13", "Gợi ý AI để con người cân nhắc", "Cung cấp tư vấn có giới hạn; CAP và con người vẫn quyết định nghiệp vụ.", "Các trường defect tối thiểu trong danh sách cho phép; không gửi thông tin bí mật hoặc tệp.", "Gợi ý đã làm sạch, độ tin cậy, giải thích và bản ghi AiSuggestions ở trạng thái PENDING.", "Người dùng được phép", "AI mặc định tắt; cô lập lỗi; kiểm tra đầu ra; không lưu prompt hoặc response thô.", "Không tự phân công, phân loại, xác nhận trùng hoặc đổi vòng đời.", "Người dùng chỉ đánh giá nội dung hiển thị và có thể thực hiện riêng một thao tác nghiệp vụ thông thường được phép; dịch vụ hiện chưa có thao tác cập nhật trạng thái đánh giá.", "AiSuggestions lưu kết quả đã chuẩn hóa ở PENDING; chưa lưu ACCEPTED, REJECTED hoặc IGNORED.", "Bản thân gợi ý không tạo thông báo."),
]

EN_DATA_OWNERSHIP = [
    ["Bugs and classification", "Tester creates; PM governs correction", "Application business data", "Retain for project audit; changes are history-backed"],
    ["Developer responsibility", "PM / authorized maintainer", "Assignment master data", "Review when team capability changes"],
    ["Comments and attachments", "Author supplies; project controls access", "Collaboration and evidence", "Attachment bytes follow configured DB/S3 profile"],
    ["HistoryEvents / HistoryLogs", "System", "Append-only audit evidence", "Do not overwrite to hide prior values"],
    ["Notifications / deliveries", "System", "Operational communication record", "Keep safe status/attempt data; never provider secrets"],
    ["AuthSessions", "System", "Security session data", "Token hash only; expire/revoke according to policy"],
    ["AiSuggestions", "System plus human reviewer", "Advisory audit record", "Normalized output only; no raw prompt/response"],
]
VI_DATA_OWNERSHIP = [
    ["Bugs và dữ liệu phân loại", "Tester tạo; PM quản trị việc hiệu chỉnh", "Dữ liệu nghiệp vụ của ứng dụng", "Giữ để kiểm tra dự án; mọi thay đổi có lịch sử"],
    ["Developer Responsibility", "PM hoặc người bảo trì được phép", "Dữ liệu nền phục vụ phân công", "Đánh giá lại khi năng lực nhóm thay đổi"],
    ["Bình luận và tệp đính kèm", "Tác giả cung cấp; dự án kiểm soát truy cập", "Trao đổi và bằng chứng", "Nội dung tệp theo cấu hình cơ sở dữ liệu/S3"],
    ["HistoryEvents / HistoryLogs", "Hệ thống", "Bằng chứng lưu vết chỉ bổ sung", "Không ghi đè để che giá trị trước đó"],
    ["Notifications / NotificationDeliveries", "Hệ thống", "Bản ghi truyền thông vận hành", "Chỉ giữ trạng thái/lần thử an toàn; không lưu bí mật nhà cung cấp"],
    ["AuthSessions", "Hệ thống", "Dữ liệu phiên bảo mật", "Chỉ giữ token băm; hết hạn hoặc thu hồi theo chính sách"],
    ["AiSuggestions", "Hệ thống", "Bản ghi tư vấn", "Chỉ giữ đầu ra đã chuẩn hóa ở PENDING; không lưu prompt/response thô"],
]

EN_INTERFACE_ROWS = [
    ["Fiori Elements / SAPUI5", "Human UI", "OData V4 metadata, list/object pages, role-aware actions", "UI visibility is not authorization"],
    ["CAP BugService", "Business API", "Entities, bound actions, validations and transactions", "Backend is authoritative"],
    ["AuthService / AuthSessions", "Authentication", "Login/logout, bearer session resolution", "Hash/token/no-secret controls"],
    ["SQLite", "Local persistence", "Developer and programmatic verification", "Not shared-QA or production acceptance"],
    ["Render / PostgreSQL", "Shared QA persistence", "CAP integration profile", "Availability and continuity evidence remain environment-dependent"],
    ["S3-compatible storage", "Attachment bytes", "Externally bound object store", "Compatibility/binding and persistence evidence required"],
    ["NotificationDeliveries", "Email boundary", "Database outbox to private provider", "At-least-once; failure isolated; local disabled is SKIPPED"],
    ["AI provider", "Advisory boundary", "Minimum allowlisted request and validated response", "Disabled by default; human review mandatory"],
]
VI_INTERFACE_ROWS = [
    ["Fiori Elements / SAPUI5", "Giao diện người dùng", "Metadata OData V4, trang danh sách/đối tượng và thao tác theo vai trò", "Hiển thị trên giao diện không thay thế phân quyền"],
    ["CAP BugService", "API nghiệp vụ", "Entity, bound action, kiểm tra dữ liệu và giao dịch", "Backend là lớp có thẩm quyền"],
    ["AuthService / AuthSessions", "Xác thực", "Đăng nhập, đăng xuất và phân giải bearer session", "Kiểm soát băm, token và không lộ bí mật"],
    ["SQLite", "Lưu trữ cục bộ", "Phát triển và xác minh bằng chương trình", "Không phải bằng chứng nghiệm thu QA dùng chung hoặc production"],
    ["Render / PostgreSQL", "Lưu trữ QA dùng chung", "CAP integration profile", "Bằng chứng sẵn sàng và liên tục phụ thuộc môi trường"],
    ["Lưu trữ tương thích S3", "Nội dung tệp đính kèm", "Kho đối tượng được liên kết ngoài", "Cần bằng chứng tương thích, liên kết và lưu bền"],
    ["NotificationDeliveries", "Ranh giới email", "Hộp thư đi trong cơ sở dữ liệu tới nhà cung cấp riêng", "Có thể giao lặp; lỗi được cô lập; local tắt là SKIPPED"],
    ["Nhà cung cấp AI", "Ranh giới tư vấn", "Yêu cầu tối thiểu theo danh sách cho phép và phản hồi đã kiểm tra", "Mặc định tắt; con người phải đánh giá"],
]

TRACE_ROWS = [
    ["BP-01", "Auth and role boundary", "BRD/SRS auth requirements; AuthService", "qa:auth:programmatic", "Executed PASS"],
    ["BP-02–BP-09", "Core defect lifecycle", "Business rules BR-01–BR-46; FRS workflows", "Current test scenarios plus three broader suites", "Partly executed; three suites Pending"],
    ["BP-10", "Comments and attachments", "SRS/FRS collaboration and storage rules", "Comments/attachments regression", "Pending in current Test Report"],
    ["BP-11", "Notification outbox", "NotificationDeliveries rules", "qa:email-outbox:programmatic", "Executed PASS; provider disabled/SKIPPED"],
    ["BP-12", "PM monitoring", "PM reporting requirements", "qa:pm-monitoring:programmatic", "Executed PASS"],
    ["BP-13", "AI human-review boundary", "AI canonical rules and AiSuggestions", "AI provider/human-review regression", "Pending in current Test Report"],
    ["Acceptance", "Mentor/UAT/sign-off", "UAT and final-report artifacts", "Mentor/user execution", "Not complete; UAT Prepared only"],
]
VI_TRACE_ROWS = [
    ["BP-01", "Ranh giới xác thực và vai trò", "Yêu cầu xác thực BRD/SRS; AuthService", "qa:auth:programmatic", "Đã chạy và PASS"],
    ["BP-02–BP-09", "Vòng đời defect cốt lõi", "Quy tắc BR-01–BR-46; luồng FRS", "Các kịch bản hiện tại và ba bộ kiểm thử mở rộng", "Đã chạy một phần; ba bộ còn Pending"],
    ["BP-10", "Bình luận và tệp đính kèm", "Quy tắc trao đổi/lưu trữ trong SRS/FRS", "Kiểm thử hồi quy bình luận/tệp", "Pending trong Test Report hiện tại"],
    ["BP-11", "Hộp thư đi thông báo", "Quy tắc NotificationDeliveries", "qa:email-outbox:programmatic", "Đã chạy và PASS; nhà cung cấp tắt/SKIPPED"],
    ["BP-12", "PM giám sát", "Yêu cầu báo cáo PM", "qa:pm-monitoring:programmatic", "Đã chạy và PASS"],
    ["BP-13", "Ranh giới AI để con người đánh giá", "Quy tắc AI chuẩn và AiSuggestions", "Kiểm thử hồi quy AI/đánh giá của con người", "Pending trong Test Report hiện tại"],
    ["Nghiệm thu", "Người hướng dẫn/UAT/ký xác nhận", "UAT và Final Project Report", "Người hướng dẫn/người dùng thực hiện", "Chưa hoàn tất; UAT mới ở trạng thái Prepared"],
]


def overview_table_specs(content: dict) -> list[tuple[str, list[str], list[list[str]], list[float]]]:
    if content is EN:
        baseline_purpose = [
            "CAP/Fiori application and OData service",
            "Authenticated identity and backend role boundary",
            "Local development and shared-QA persistence",
            "Store and retrieve defect evidence",
            "In-app events and asynchronous email attempts",
            "Optional review-only assistance",
        ]
        baseline_limits = [
            "Educational/shared-QA baseline; no production certification",
            "Custom bearer sessions; not enterprise SSO",
            "Continuity remains environment-dependent",
            "Shared-QA acceptance needs active S3 evidence",
            "Live-provider success is not claimed",
            "Disabled by default; review state remains PENDING",
        ]
        capability_roles = ["Tester", "Tester / reviewer", "Tester / PM", "Tester / PM", "All business roles", "All business roles", "System", "PM", "Authorized user"]
        capability_processes = ["BP-02", "BP-03", "BP-04", "BP-05", "BP-06–BP-09", "BP-10", "BP-11", "BP-12", "BP-13"]
        object_sources = ["Tester / PM", "System / PM", "PM / reference data", "PM", "Authorized users", "System", "System", "Human decision / System"]
        object_uses = ["All business roles", "Authentication and authorization", "Classification and assignment", "Assignment", "Collaboration and QA", "Audit", "Recipients and operations", "Duplicate/AI review"]
        object_methods = ["CAP entity / OData V4", "AuthService and bearer session", "Associations and value help", "CAP responsibility lookup", "Composition; DB/S3 attachments", "Append-only compositions", "Database outbox / worker", "Confirmed link / read-only audit"]
        integration_owners = ["Developer environment", "Platform configuration", "System / operations", "System / authorized user"]
        integration_uses = ["Local verification", "Shared QA", "Notification recipients", "Advisory actions"]
        integration_methods = ["CAP SQLite adapter", "HTTPS/OData V4; PostgreSQL/S3 bindings", "NotificationDeliveries worker / SMTP", "Server-side allowlisted provider request"]
        ownership_storage = ["PostgreSQL / SQLite", "Relational master data", "Database and configured S3", "Relational append-only records", "Database outbox", "Database token hashes", "Database normalized audit"]
        ownership_security = ["Role-filtered access", "PM-maintained", "Authorized readers only", "Do not overwrite prior evidence", "No provider secrets", "No raw password or token", "No raw prompt/response; PENDING only"]
        interface_callers = ["Browser user", "Fiori/Auth clients", "Browser login", "CAP runtime", "CAP runtime", "Attachment service", "CAP event / worker", "CAP AI action"]
        interface_receivers = ["CAP BugService", "CAP services", "AuthService / AuthSessions", "SQLite", "PostgreSQL", "S3 object store", "Email provider", "AI provider"]
        titles = ["Current solution baseline", "Functional capabilities", "Information objects and integrations", "Data ownership and retention", "Interfaces and control boundaries"]
        headers = [
            ["Layer", "Current implementation", "Purpose", "Current limitation"],
            ["Capability ID", "Capability", "Primary role", "Main result", "Related process"],
            ["Object / System", "Stored information", "Source / Owner", "Used by", "Integration method"],
            ["Data object", "Business owner", "Technical storage", "Retention / current handling", "Security note"],
            ["Boundary", "Caller", "Receiver", "Protocol / Interface", "Validation / Control"],
        ]
    else:
        baseline_purpose = [
            "Ứng dụng CAP/Fiori và dịch vụ OData",
            "Xác lập danh tính và ranh giới vai trò ở backend",
            "Lưu trữ cho phát triển cục bộ và QA dùng chung",
            "Lưu và truy xuất bằng chứng defect",
            "Ghi sự kiện trong ứng dụng và lần thử email bất đồng bộ",
            "Hỗ trợ tùy chọn để con người đánh giá",
        ]
        baseline_limits = [
            "Nền giáo dục/QA dùng chung; chưa chứng nhận production",
            "Phiên bearer tùy chỉnh; không phải SSO doanh nghiệp",
            "Tính liên tục phụ thuộc môi trường",
            "Nghiệm thu QA cần bằng chứng S3 đang hoạt động",
            "Chưa tuyên bố thành công qua nhà cung cấp thực",
            "Mặc định tắt; trạng thái đánh giá vẫn là PENDING",
        ]
        capability_roles = ["Tester", "Tester / người đánh giá", "Tester / PM", "Tester / PM", "Các vai trò nghiệp vụ", "Các vai trò nghiệp vụ", "Hệ thống", "PM", "Người dùng được phép"]
        capability_processes = ["BP-02", "BP-03", "BP-04", "BP-05", "BP-06–BP-09", "BP-10", "BP-11", "BP-12", "BP-13"]
        object_sources = ["Tester / PM", "Hệ thống / PM", "PM / dữ liệu tham chiếu", "PM", "Người dùng được phép", "Hệ thống", "Hệ thống", "Quyết định của con người / Hệ thống"]
        object_uses = ["Các vai trò nghiệp vụ", "Xác thực và phân quyền", "Phân loại và phân công", "Phân công", "Cộng tác và QA", "Kiểm toán", "Người nhận và vận hành", "Đánh giá trùng lặp/AI"]
        object_methods = ["CAP entity / OData V4", "AuthService và bearer session", "Association và value help", "Tra cứu trách nhiệm trong CAP", "Composition; tệp trong DB/S3", "Composition chỉ bổ sung", "Hộp thư đi / worker", "Liên kết đã xác nhận / audit chỉ đọc"]
        integration_owners = ["Môi trường Developer", "Cấu hình nền tảng", "Hệ thống / vận hành", "Hệ thống / người dùng được phép"]
        integration_uses = ["Xác minh cục bộ", "QA dùng chung", "Người nhận thông báo", "Thao tác AI tư vấn"]
        integration_methods = ["CAP SQLite adapter", "HTTPS/OData V4; liên kết PostgreSQL/S3", "NotificationDeliveries worker / SMTP", "Yêu cầu phía máy chủ theo danh sách cho phép"]
        ownership_storage = ["PostgreSQL / SQLite", "Dữ liệu chủ quan hệ", "Cơ sở dữ liệu và S3 đã cấu hình", "Bản ghi quan hệ chỉ bổ sung", "Hộp thư đi trong cơ sở dữ liệu", "Token băm trong cơ sở dữ liệu", "Audit đã chuẩn hóa trong cơ sở dữ liệu"]
        ownership_security = ["Truy cập được lọc theo vai trò", "PM bảo trì", "Chỉ người đọc được phép", "Không ghi đè bằng chứng cũ", "Không lưu bí mật nhà cung cấp", "Không lưu mật khẩu hoặc token thô", "Không lưu prompt/response thô; chỉ PENDING"]
        interface_callers = ["Người dùng trình duyệt", "Fiori/Auth client", "Màn hình đăng nhập", "CAP runtime", "CAP runtime", "Dịch vụ tệp đính kèm", "Sự kiện CAP / worker", "Thao tác AI của CAP"]
        interface_receivers = ["CAP BugService", "Các dịch vụ CAP", "AuthService / AuthSessions", "SQLite", "PostgreSQL", "Kho đối tượng S3", "Nhà cung cấp email", "Nhà cung cấp AI"]
        titles = ["Hiện trạng giải pháp", "Các năng lực chức năng", "Đối tượng thông tin và tích hợp", "Quyền sở hữu và lưu giữ dữ liệu", "Giao diện và ranh giới kiểm soát"]
        headers = [
            ["Lớp", "Hiện trạng triển khai", "Mục đích", "Giới hạn hiện tại"],
            ["Mã năng lực", "Năng lực", "Vai trò chính", "Kết quả chính", "Quy trình liên quan"],
            ["Đối tượng / Hệ thống", "Thông tin lưu trữ", "Nguồn / Chủ sở hữu", "Bên sử dụng", "Phương thức tích hợp"],
            ["Đối tượng dữ liệu", "Chủ sở hữu nghiệp vụ", "Lưu trữ kỹ thuật", "Lưu giữ / xử lý hiện tại", "Lưu ý bảo mật"],
            ["Ranh giới", "Bên gọi", "Bên nhận", "Giao thức / Giao diện", "Kiểm tra / Kiểm soát"],
        ]

    baseline_rows = [row + [baseline_purpose[index], baseline_limits[index]] for index, row in enumerate(content["baseline_rows"])]
    capability_rows = [[row[0], f"{row[1]} ({row[3]})", capability_roles[index], row[2], capability_processes[index]] for index, row in enumerate(content["capability_rows"])]
    object_rows = [row + [object_sources[index], object_uses[index], object_methods[index]] for index, row in enumerate(content["entity_rows"])]
    for index, row in enumerate(content["integration_rows"]):
        object_rows.append([row[0], f"{row[1]}; {row[2]}", integration_owners[index], integration_uses[index], integration_methods[index]])
    ownership_source = EN_DATA_OWNERSHIP if content is EN else VI_DATA_OWNERSHIP
    ownership_rows = [[row[0], row[1], ownership_storage[index], row[3], ownership_security[index]] for index, row in enumerate(ownership_source)]
    interface_source = EN_INTERFACE_ROWS if content is EN else VI_INTERFACE_ROWS
    interface_rows = [[row[0], interface_callers[index], interface_receivers[index], row[2], row[3]] for index, row in enumerate(interface_source)]
    rows = [baseline_rows, capability_rows, object_rows, ownership_rows, interface_rows]
    widths = [[1.5, 2.0, 1.8, 1.8], [1.1, 1.5, 1.15, 1.9, 1.45], [1.65, 1.8, 1.15, 1.0, 1.5], [1.2, 1.2, 1.4, 1.9, 1.4], [1.2, 1.0, 1.0, 1.6, 2.3]]
    return list(zip(titles, headers, rows, widths, strict=True))


def organization_table_specs(content: dict) -> list[tuple[str, list[str], list[list[str]], list[float]]]:
    if content is EN:
        allowed = [
            "Create/classify; supply information; retest; close/reopen",
            "Review assigned work; request information; reject; resolve",
            "Monitor and coordinate queues, assignment and escalation",
            "Validate, persist, audit and deliver configured events",
        ]
        restricted = [
            "Cannot bypass validation or act as an unassigned Developer",
            "Cannot process unassigned work or bypass transition rules",
            "Cannot become a second assignee or bypass backend controls",
            "No autonomous business or AI decision",
        ]
        titles = ["Roles and responsibilities", "RACI responsibility matrix"]
        role_headers = ["Role", "Main responsibilities", "Allowed actions", "Restricted actions"]
    else:
        allowed = [
            "Tạo/phân loại; bổ sung thông tin; kiểm thử lại; đóng/mở lại",
            "Đánh giá việc được phân công; yêu cầu thông tin; từ chối; xử lý",
            "Giám sát và điều phối hàng đợi, phân công và cảnh báo",
            "Kiểm tra, lưu, ghi nhật ký và giao sự kiện đã cấu hình",
        ]
        restricted = [
            "Không bỏ qua validation hoặc xử lý như Developer chưa được phân công",
            "Không xử lý việc chưa được phân công hoặc bỏ qua quy tắc chuyển trạng thái",
            "Không trở thành assignee thứ hai hoặc bỏ qua kiểm soát backend",
            "Không tự quyết định nghiệp vụ hoặc quyết định AI",
        ]
        titles = ["Vai trò và trách nhiệm", "Ma trận trách nhiệm RACI"]
        role_headers = ["Vai trò", "Trách nhiệm chính", "Thao tác được phép", "Thao tác bị hạn chế"]
    role_rows = [[row[0], row[1], allowed[index], restricted[index]] for index, row in enumerate(content["role_rows"])]
    return [
        (titles[0], role_headers, role_rows, [1.0, 2.4, 2.1, 1.6]),
        (titles[1], ["Activity" if content is EN else "Hoạt động", "Tester", "Developer", "PM", "System" if content is EN else "Hệ thống"], content["raci_rows"], [2.7, 1.0, 1.25, 0.9, 1.25]),
    ]


def lifecycle_table_spec(content: dict) -> tuple[str, list[str], list[list[str]], list[float]]:
    if content is EN:
        rows = [
            ["Create", "Submit valid defect", "Tester / PM", "Assigned or Pending Assignment", "Set valid Developer or remain empty", "Developer or Tester/PM queue"],
            ["New", "Legacy/import handling only", "System", "No normal-flow claim", "Preserve imported value", "As imported"],
            ["Pending Assignment", "Assign valid Developer", "Tester / PM", "Assigned", "Set Developer", "Developer"],
            ["Assigned", "Start technical review", "Assigned Developer", "In Review / In Progress", "Unchanged", "Developer"],
            ["Assigned / In Review / In Progress", "Request more information", "Assigned Developer", "Need More Information", "Unchanged", "Tester / PM"],
            ["Need More Information", "Resubmit information", "Tester / PM", "Assigned", "Keep existing Developer", "Developer"],
            ["Assigned / In Review / In Progress", "Reject with reason", "Assigned Developer", "Rejected", "Preserve until correction", "Tester / PM"],
            ["Rejected", "Correct and assign/reassign", "Tester / PM", "Assigned or Pending Assignment", "Set valid Developer or clear", "Developer or Tester/PM queue"],
            ["In Review / In Progress", "Resolve with required context", "Assigned Developer", "Resolved", "Unchanged", "Tester / PM"],
            ["Resolved", "Retest / accept / reopen", "Tester / PM", "Retest Required / Closed / Reopened", "Follow authorized flow", "Tester / PM or Developer"],
            ["Retest Required", "Record pass or failure", "Tester / PM", "Closed or Reopened", "Follow authorized flow", "Tester / PM or Developer"],
            ["Closed", "Reopen with reason", "Tester / PM", "Reopened", "Follow authorized assignment flow", "Tester / PM or Developer"],
        ]
        return "Bug lifecycle, status transition and next processor", ["Current status", "User action", "Allowed role", "Next status", "Assignee effect", "Next processor"], rows, [1.25, 1.2, 1.1, 1.2, 1.15, 1.2]
    rows = [
        ["Tạo mới", "Gửi defect hợp lệ", "Tester / PM", "Assigned hoặc Pending Assignment", "Gán Developer hợp lệ hoặc để trống", "Developer hoặc hàng đợi Tester/PM"],
        ["New", "Chỉ xử lý dữ liệu cũ/nhập", "Hệ thống", "Không tuyên bố luồng thường", "Giữ giá trị đã nhập", "Theo dữ liệu nhập"],
        ["Pending Assignment", "Phân công Developer hợp lệ", "Tester / PM", "Assigned", "Gán Developer", "Developer"],
        ["Assigned", "Bắt đầu đánh giá kỹ thuật", "Developer được phân công", "In Review / In Progress", "Không đổi", "Developer"],
        ["Assigned / In Review / In Progress", "Yêu cầu thêm thông tin", "Developer được phân công", "Need More Information", "Không đổi", "Tester / PM"],
        ["Need More Information", "Gửi lại thông tin", "Tester / PM", "Assigned", "Giữ Developer hiện hữu", "Developer"],
        ["Assigned / In Review / In Progress", "Từ chối kèm lý do", "Developer được phân công", "Rejected", "Giữ tới khi hiệu chỉnh", "Tester / PM"],
        ["Rejected", "Hiệu chỉnh và phân công lại", "Tester / PM", "Assigned hoặc Pending Assignment", "Gán Developer hợp lệ hoặc xóa", "Developer hoặc hàng đợi Tester/PM"],
        ["In Review / In Progress", "Xác nhận xử lý với thông tin bắt buộc", "Developer được phân công", "Resolved", "Không đổi", "Tester / PM"],
        ["Resolved", "Kiểm thử lại / chấp nhận / mở lại", "Tester / PM", "Retest Required / Closed / Reopened", "Theo luồng được phép", "Tester / PM hoặc Developer"],
        ["Retest Required", "Ghi kết quả đạt hoặc không đạt", "Tester / PM", "Closed hoặc Reopened", "Theo luồng được phép", "Tester / PM hoặc Developer"],
        ["Closed", "Mở lại kèm lý do", "Tester / PM", "Reopened", "Theo luồng phân công được phép", "Tester / PM hoặc Developer"],
    ]
    return "Vòng đời bug, chuyển trạng thái và người xử lý tiếp", ["Trạng thái hiện tại", "Thao tác người dùng", "Vai trò được phép", "Trạng thái tiếp theo", "Ảnh hưởng assignee", "Người xử lý tiếp"], rows, [1.25, 1.2, 1.1, 1.2, 1.15, 1.2]


def report_table_specs(content: dict) -> list[tuple[str, list[str], list[list[str]], list[float]]]:
    source = TRACE_ROWS if content is EN else VI_TRACE_ROWS
    if content is EN:
        gaps = ["None in current local run", "Three broader suites remain Pending", "Current report still Pending", "Live provider disabled / SKIPPED", "None in current local run", "Provider/human-review test Pending", "UAT execution and sign-off Pending"]
        limitation_impacts = ["No formal acceptance", "No user acceptance result", "Email acceptance incomplete", "Attachment acceptance incomplete", "AI acceptance incomplete", "No production-readiness claim"]
        mitigations = ["Keep Draft/Pending labels", "Keep UAT Prepared", "Use in-app/outbox evidence only", "Require binding and persistence proof", "Keep disabled; disclose PENDING-only audit", "Limit claims to educational/shared QA"]
        owners = ["Mentor / DonHV", "Mentor / users", "Operations / DonHV", "Operations / QA", "DonHV / QA", "PM / architecture"]
        titles = ["Verification and acceptance status", "Known limitations"]
        headers = [["Area", "Verification method", "Current evidence", "Status", "Remaining gap"], ["ID", "Limitation", "Impact", "Current mitigation", "Follow-up / Owner"]]
    else:
        gaps = ["Không còn khoảng trống trong lần chạy local hiện tại", "Ba bộ kiểm thử mở rộng còn Pending", "Báo cáo hiện tại vẫn Pending", "Nhà cung cấp thực tắt / SKIPPED", "Không còn khoảng trống trong lần chạy local hiện tại", "Kiểm thử nhà cung cấp/người đánh giá còn Pending", "UAT và ký xác nhận còn Pending"]
        limitation_impacts = ["Chưa có nghiệm thu chính thức", "Chưa có kết quả chấp nhận của người dùng", "Nghiệm thu email chưa đầy đủ", "Nghiệm thu tệp chưa đầy đủ", "Nghiệm thu AI chưa đầy đủ", "Không được tuyên bố sẵn sàng production"]
        mitigations = ["Giữ nhãn Bản nháp/Đang chờ", "Giữ UAT ở trạng thái Prepared", "Chỉ dùng bằng chứng trong ứng dụng/hộp thư đi", "Yêu cầu bằng chứng liên kết và lưu bền", "Giữ mặc định tắt; công bố audit chỉ PENDING", "Giới hạn tuyên bố ở nền giáo dục/QA dùng chung"]
        owners = ["Người hướng dẫn / DonHV", "Người hướng dẫn / người dùng", "Vận hành / DonHV", "Vận hành / QA", "DonHV / QA", "PM / kiến trúc"]
        titles = ["Trạng thái kiểm chứng và nghiệm thu", "Các giới hạn đã biết"]
        headers = [["Phạm vi", "Phương pháp kiểm chứng", "Bằng chứng hiện tại", "Trạng thái", "Khoảng trống còn lại"], ["Mã", "Giới hạn", "Ảnh hưởng", "Biện pháp hiện tại", "Theo dõi / Chủ sở hữu"]]
    verification_rows = [[row[1], row[3], row[2], row[4], gaps[index]] for index, row in enumerate(source)]
    limitation_rows = [[f"LIM-{index + 1:02d}", item, limitation_impacts[index], mitigations[index], owners[index]] for index, item in enumerate(content["limitation_items"])]
    return [
        (titles[0], headers[0], verification_rows, [1.2, 1.6, 1.6, 1.0, 1.7]),
        (titles[1], headers[1], limitation_rows, [0.55, 2.0, 1.3, 1.9, 1.35]),
    ]


def copy_element_property(target, source, property_name: str) -> None:
    target_property = getattr(target, property_name, None)
    if target_property is not None:
        target.remove(target_property)
    source_property = getattr(source, property_name, None)
    if source_property is not None:
        target.insert(0, deepcopy(source_property))


def add_wrap_opportunities(text: str) -> str:
    """Add invisible breaks only at slash or CamelCase/entity boundaries."""
    value = str(text)
    value = re.sub(r"/(?=\S)", "/\u200b", value)

    def polish_token(match: re.Match[str]) -> str:
        token = match.group(0)
        if len(token) < 12:
            return token
        return re.sub(r"(?<=[a-z])(?=[A-Z])|(?<=[A-Z])(?=[A-Z][a-z])", "\u200b", token)

    return re.sub(r"[A-Za-z][A-Za-z0-9]*", polish_token, value)


def fill_new_table_cell(cell, text: str, prototype_cell) -> None:
    """Fill a new cell while inheriting the official Reports-table formatting."""
    copy_element_property(cell._tc, prototype_cell._tc, "tcPr")
    paragraph = cell.paragraphs[0]
    prototype_paragraph = prototype_cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        cell._tc.remove(extra._p)
    for child in list(paragraph._p):
        if child.tag in {qn("w:r"), qn("w:hyperlink"), qn("w:pPr")}:
            paragraph._p.remove(child)
    if prototype_paragraph._p.pPr is not None:
        paragraph._p.insert(0, deepcopy(prototype_paragraph._p.pPr))
    run = paragraph.add_run(add_wrap_opportunities(text))
    prototype_run = prototype_paragraph.runs[0] if prototype_paragraph.runs else None
    if prototype_run is not None and prototype_run._r.rPr is not None:
        run._r.insert(0, deepcopy(prototype_run._r.rPr))


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, prototype_table=None):
    table = doc.add_table(rows=1, cols=len(headers))
    if prototype_table is None:
        raise ValueError("Approved Blueprint content tables require an official-template prototype table")
    copy_element_property(table._tbl, prototype_table._tbl, "tblPr")
    for index, header in enumerate(headers):
        prototype_cell = prototype_table.rows[0].cells[min(index, len(prototype_table.columns) - 1)]
        fill_new_table_cell(table.rows[0].cells[index], header, prototype_cell)
    keep_row_together(table.rows[0], repeat_header=True)
    for values in rows:
        cells = table.add_row().cells
        for index in range(len(headers)):
            prototype_column = 0 if index == 0 else min(index, len(prototype_table.columns) - 1)
            fill_new_table_cell(cells[index], values[index] if index < len(values) else "", prototype_table.rows[1].cells[prototype_column])
        keep_row_together(table.rows[-1])
    if widths:
        set_table_grid_widths(table, widths)
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_process_details(doc: Document, content: dict) -> None:
    details = EN_PROCESS_DETAILS if content is EN else VI_PROCESS_DETAILS
    labels = (
        ["Objective", "Inputs", "Outputs", "Primary owner", "Validations / controls", "Status effect", "Next action owner", "History effect", "Notification effect"]
        if content is EN
        else ["Mục tiêu", "Input", "Output", "Owner chính", "Validation / control", "Status effect", "Owner action tiếp theo", "History effect", "Notification effect"]
    )
    for index, detail in enumerate(details):
        process_id, name, *values = detail
        heading = doc.add_heading(f"{process_id} — {name}", level=3)
        heading.paragraph_format.keep_with_next = True
        rows = [[label, value] for label, value in zip(labels, values, strict=True)]
        add_table(doc, ["Attribute" if content is EN else "Thuộc tính", "Blueprint definition" if content is EN else "Định nghĩa Blueprint"], rows, [1.45, 5.55])


def add_embedded_diagram(doc: Document, asset: str, caption: str, page_break_before: bool = False) -> None:
    image_path = ROOT / "docs" / "diagrams" / "rendered" / "png" / f"{asset}.png"
    if not image_path.exists() or image_path.stat().st_size == 0:
        raise FileNotFoundError(f"Missing Blueprint diagram image: {image_path}")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.page_break_before = page_break_before
    run = paragraph.add_run()
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    scale = min(6.35 / width_px, 8.15 / height_px)
    run.add_picture(
        str(image_path),
        width=Inches(width_px * scale),
        height=Inches(height_px * scale),
    )
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.italic = True
    caption_run.font.size = Pt(9)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def add_toc(paragraph, cached_text: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    begin_run._r.extend([begin, instruction])

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    cached_run = paragraph.add_run(cached_text)
    cached_run.font.name = "Arial"
    cached_run.font.size = Pt(10.5)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.insert_element_before(
            existing,
            "w:hdrShapeDefaults",
            "w:footnotePr",
            "w:endnotePr",
            "w:compat",
            "w:docVars",
            "w:rsids",
            "m:mathPr",
            "w:uiCompat97To2003",
            "w:attachedSchema",
            "w:themeFontLang",
            "w:clrSchemeMapping",
            "w:doNotIncludeSubdocsInStats",
            "w:doNotAutoCompressPictures",
            "w:forceUpgrade",
            "w:captions",
            "w:readModeInkLockDown",
            "w:smartTagType",
            "sl:schemaLibrary",
            "w:shapeDefaults",
            "w:doNotEmbedSmartTags",
            "w:decimalSymbol",
            "w:listSeparator",
        )
    existing.set(qn("w:val"), "1")


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for level, size in [(1, 16), (2, 13), (3, 11)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def configure_section(doc: Document, footer_text: str, page_label: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("Issue and Defect Tracking System in SAP")
    header_run.font.name = "Arial"
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = RGBColor(89, 89, 89)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"{footer_text} | {page_label} ")
    add_page_field(footer)


def add_cover(doc: Document, content: dict) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(90)
    run = title.add_run(content["title"])
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(content["subtitle"])
    subtitle_run.bold = True
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(16)
    cover_values = (
        [f"Project / Group: IDTS SAP01", f"Language: {content['language']}", f"Version: {VERSION}", f"Status: {content['status']}", f"Prepared by: {content['prepared']}", f"Date: {DATE}"]
        if content is EN
        else [f"Dự án / Nhóm: IDTS SAP01", f"Ngôn ngữ: {content['language']}", f"Phiên bản: {VERSION}", f"Trạng thái: {content['status']}", f"Người chuẩn bị: {content['prepared']}", f"Ngày: {DATE}"]
    )
    for value in cover_values:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(value)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run("SAP490 mentor-review artifact | Repository sources remain canonical")
    note_run.italic = True
    note_run.font.color.rgb = RGBColor(89, 89, 89)




def fill_template_cell(cell, text: str) -> None:
    """Replace cell text while retaining template cell/paragraph properties."""
    paragraph = cell.paragraphs[0]
    run_properties = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run_properties = deepcopy(paragraph.runs[0]._r.rPr)
    for extra in cell.paragraphs[1:]:
        cell._tc.remove(extra._p)
    for child in list(paragraph._p):
        if child.tag in {qn("w:r"), qn("w:hyperlink")}:
            paragraph._p.remove(child)
    run = paragraph.add_run(str(text))
    if run_properties is not None:
        if run._r.rPr is not None:
            run._r.remove(run._r.rPr)
        run._r.insert(0, run_properties)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def ensure_rows(table, count: int) -> None:
    while len(table.rows) < count:
        table._tbl.append(deepcopy(table.rows[-1]._tr))


def resize_rows(table, count: int) -> None:
    ensure_rows(table, count)
    while len(table.rows) > count:
        table._tbl.remove(table.rows[-1]._tr)


def keep_row_together(row, repeat_header: bool = False) -> None:
    row_properties = row._tr.get_or_add_trPr()
    if row_properties.find(qn("w:cantSplit")) is None:
        row_properties.append(OxmlElement("w:cantSplit"))
    if repeat_header and row_properties.find(qn("w:tblHeader")) is None:
        row_properties.append(OxmlElement("w:tblHeader"))


def replace_part_text(part, replacements: dict[str, str]) -> None:
    for node in part._element.xpath(".//w:t"):
        if node.text in replacements:
            node.text = replacements[node.text]


def remove_sample_objects(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in list(paragraph._p):
            if next(run.iter(qn("w:object")), None) is not None or next(run.iter(qn("w:drawing")), None) is not None:
                paragraph._p.remove(run)


def set_template_paragraph(paragraph, text: str, style: str | None = None) -> None:
    paragraph.clear()
    if style:
        paragraph.style = style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run(text)


def set_cell_width(cell, width: float) -> None:
    """Set a stable table-grid width without replacing the template table."""
    cell.width = Inches(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.first_child_found_in("w:tcW")
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:w"), str(int(width * 1440)))
    tc_width.set(qn("w:type"), "dxa")


def set_table_grid_widths(table, widths: list[float]) -> None:
    """Update both the table grid and cell preferences for cross-renderer widths."""
    table.autofit = False
    grid_columns = table._tbl.tblGrid.gridCol_lst
    for column_index, width in enumerate(widths):
        table.columns[column_index].width = Inches(width)
        if column_index < len(grid_columns):
            grid_columns[column_index].set(qn("w:w"), str(int(width * 1440)))
        for cell in table.columns[column_index].cells:
            set_cell_width(cell, width)


def block_element(block):
    return block._p if hasattr(block, "_p") else block._tbl


def set_table_title(paragraph: Paragraph, text: str, heading_level: int) -> None:
    set_template_paragraph(paragraph, text, f"Heading {heading_level}")
    paragraph.paragraph_format.keep_with_next = True


def add_table_series_after(
    doc: Document,
    anchor: Paragraph,
    specs: list[tuple[str, list[str], list[list[str]], list[float]]],
    prototype_table,
    heading_level: int,
):
    """Place approved content tables after an existing template paragraph."""
    current = anchor
    for index, (title, headers, rows, widths) in enumerate(specs):
        if index == 0:
            title_paragraph = anchor
        else:
            title_paragraph = doc.add_paragraph()
            block_element(current).addnext(title_paragraph._p)
        set_table_title(title_paragraph, title, heading_level)
        table = add_table(doc, headers, rows, widths, prototype_table)
        title_paragraph._p.addnext(table._tbl)
        current = table
    return current


def insert_picture_after(doc: Document, anchor, asset: str, caption: str, width: float = 5.8) -> Paragraph:
    image_path = ROOT / "docs" / "diagrams" / "rendered" / "png" / f"{asset}.png"
    if not image_path.exists() or image_path.stat().st_size == 0:
        raise FileNotFoundError(f"Missing Blueprint diagram image: {image_path}")
    picture_paragraph = doc.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    block_element(anchor).addnext(picture_paragraph._p)
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.italic = True
    caption_run.font.size = Pt(8)
    picture_paragraph._p.addnext(caption_paragraph._p)
    return caption_paragraph


def rows_to_lines(rows: list[list[str]], separator: str = " — ") -> str:
    return "\n".join(separator.join(str(value) for value in row if str(value).strip()) for row in rows)


def fill_cover(doc: Document, content: dict) -> None:
    title_table, info_table = doc.tables[0].cell(1, 0).tables
    fill_template_cell(title_table.cell(0, 0), "BUSINESS BLUEPRINT")
    labels = (
        ["Document information", "Project Name", "Module", "Created by", "Version", "Date"]
        if content is EN
        else ["Thông tin tài liệu", "Tên dự án", "Phân hệ", "Người tạo", "Phiên bản", "Ngày"]
    )
    values = ["", content["title"], "IDTS SAP01", "DonHV / IDTS SAP01 Team", VERSION, DATE]
    for row_index, label in enumerate(labels):
        fill_template_cell(info_table.cell(row_index, 0), label)
        fill_template_cell(info_table.cell(row_index, 1), values[row_index])


def fill_control_tables(doc: Document, content: dict) -> None:
    history, advisor, fu = doc.tables[1], doc.tables[2], doc.tables[3]
    headers = (
        ["Change history", "Changed date", "Items changed", "Changed content / reason", "Updated by", "Type (A/C/D)", "Version"]
        if content is EN
        else ["Lịch sử thay đổi", "Ngày thay đổi", "Hạng mục thay đổi", "Nội dung / lý do", "Người cập nhật", "Loại (A/C/D)", "Phiên bản"]
    )
    fill_template_cell(history.cell(0, 0), headers[0])
    for index, label in enumerate(headers[1:]):
        fill_template_cell(history.cell(1, index), label)
    history_rows = (
        [
            [DATE, "Official-template baseline", "Consolidated IDTS content in the official SAP490 template", "DonHV", "C", "v0.3"],
            [DATE, "Table-layout remediation", "Converted ten approved content sections to true Word tables; business meaning unchanged", "DonHV", "C", VERSION],
        ]
        if content is EN
        else [
            [DATE, "Bản chuẩn theo template chính thức", "Tổng hợp nội dung IDTS trong template SAP490 chính thức", "DonHV", "C", "v0.3"],
            [DATE, "Khắc phục cách trình bày bằng bảng", "Chuyển mười phần nội dung đã duyệt sang bảng Word thật; không thay đổi ý nghĩa nghiệp vụ", "DonHV", "C", VERSION],
        ]
    )
    for row_index, values in enumerate(history_rows, 2):
        for column_index, value in enumerate(values):
            fill_template_cell(history.cell(row_index, column_index), value)

    if content is not EN:
        fill_template_cell(advisor.cell(0, 0), "Xác nhận của người hướng dẫn")
        fill_template_cell(fu.cell(0, 0), "Xác nhận tiếp theo")
        for table in (advisor, fu):
            for column_index, value in enumerate(["", "Họ tên và vai trò", "Chữ ký", "Ngày", "Ghi chú"]):
                fill_template_cell(table.cell(1, column_index), value)

    labels = ["Created by", "Reviewed by", "Approved by"] if content is EN else ["Người tạo", "Người đánh giá", "Người phê duyệt"]
    values = (
        [
            ["DonHV / Project Leader", "", DATE, "Prepared"],
            ["Mentor / Supervisor", "", "", "Pending review"],
            ["Mentor / Supervisor", "", "", "Pending approval"],
        ]
        if content is EN
        else [
            ["DonHV / Trưởng dự án", "", DATE, "Đã chuẩn bị"],
            ["Mentor / Supervisor", "", "", "Chờ đánh giá"],
            ["Mentor / Supervisor", "", "", "Chờ phê duyệt"],
        ]
    )
    for row_index in range(2, 5):
        fill_template_cell(advisor.cell(row_index, 0), labels[row_index - 2])
        for column_index, value in enumerate(values[row_index - 2], 1):
            fill_template_cell(advisor.cell(row_index, column_index), value)
    labels = ["Reviewed by", "Approved by"] if content is EN else ["Người đánh giá", "Người phê duyệt"]
    for row_index in range(2, 4):
        fill_template_cell(fu.cell(row_index, 0), labels[row_index - 2])
        fill_template_cell(fu.cell(row_index, 1), "Mentor / Supervisor")
        fill_template_cell(fu.cell(row_index, 2), "")
        fill_template_cell(fu.cell(row_index, 3), "")
        state = ("Pending review" if row_index == 2 else "Pending approval") if content is EN else ("Chờ đánh giá" if row_index == 2 else "Chờ phê duyệt")
        fill_template_cell(fu.cell(row_index, 4), state)


def fill_glossary(doc: Document, content: dict) -> None:
    table = doc.tables[5]
    headers = ["Term", "Definition", "Note"] if content is EN else ["Thuật ngữ", "Định nghĩa", "Ghi chú"]
    for column_index, value in enumerate(headers):
        fill_template_cell(table.cell(0, column_index), value)
    additional = (
        [
            ["AuthSessions", "Server-managed authenticated session records", "No raw password storage"],
            ["NotificationDeliveries", "Asynchronous email delivery outbox", "Provider failure cannot roll back workflow"],
            ["nextProcessor", "Current person or queue expected to act", "Does not replace the technical assignee"],
            ["AiSuggestions", "Normalized advisory-AI review record", "Human review is mandatory"],
        ]
        if content is EN
        else [
            ["AuthSessions", "Bản ghi phiên xác thực do máy chủ quản lý", "Không lưu mật khẩu thô"],
            ["NotificationDeliveries", "Hộp thư đi phục vụ giao email bất đồng bộ", "Lỗi nhà cung cấp không hoàn tác quy trình"],
            ["nextProcessor", "Người hoặc hàng đợi cần thực hiện thao tác tiếp theo", "Không thay thế assignee kỹ thuật"],
            ["AiSuggestions", "Bản ghi AI tư vấn đã chuẩn hóa", "Hiện chỉ lưu trạng thái PENDING"],
        ]
    )
    rows = [list(row) + [""] if len(row) == 2 else list(row) for row in content["glossary_rows"]]
    existing_terms = {row[0].casefold() for row in rows}
    rows.extend(row for row in additional if row[0].casefold() not in existing_terms)
    resize_rows(table, len(rows) + 1)
    for row_index, values in enumerate(rows, 1):
        for column_index in range(3):
            fill_template_cell(table.cell(row_index, column_index), values[column_index] if column_index < len(values) else "")


def fill_process_table(doc: Document, content: dict) -> None:
    table = doc.tables[6]
    column_widths = [0.9, 1.65, 3.55, 1.1]
    headers = ["Step #", "Step Name", "Detailed Description", "Role"] if content is EN else ["Bước", "Tên bước", "Mô tả chi tiết", "Vai trò"]
    for column_index, value in enumerate(headers):
        fill_template_cell(table.cell(0, column_index), value)
    details = EN_PROCESS_DETAILS if content is EN else VI_PROCESS_DETAILS
    resize_rows(table, len(details) + 1)
    set_table_grid_widths(table, column_widths)
    keep_row_together(table.rows[0], repeat_header=True)
    for row_index, detail in enumerate(details, 1):
        process_id, name, objective, inputs, outputs, owner, controls, status_effect, next_owner, history_effect, notification_effect = detail
        if content is EN:
            description = "\n".join([
                f"Purpose: {objective}",
                f"Input/output: {inputs} → {outputs}",
                f"Control: {controls}",
                f"State/next owner: {status_effect} | {next_owner}",
                f"Audit/notification: {history_effect} | {notification_effect}",
            ])
        else:
            description = "\n".join([
                f"Mục đích: {objective}",
                f"Đầu vào/đầu ra: {inputs} → {outputs}",
                f"Kiểm soát: {controls}",
                f"Trạng thái/người xử lý tiếp: {status_effect} | {next_owner}",
                f"Nhật ký/thông báo: {history_effect} | {notification_effect}",
            ])
        display_process_id = process_id.replace("-", "‑")
        for column_index, value in enumerate([display_process_id, name, description, owner]):
            fill_template_cell(table.cell(row_index, column_index), value)
        keep_row_together(table.rows[row_index])


def fill_reports_table(doc: Document, content: dict) -> None:
    table = doc.tables[7]
    headers = ["No.", "Description", "Audience"] if content is EN else ["STT", "Mô tả", "Đối tượng"]
    for column_index, value in enumerate(headers):
        fill_template_cell(table.cell(0, column_index), value)
    resize_rows(table, len(content["report_rows"]) + 1)
    for row_index, values in enumerate(content["report_rows"], 1):
        row_values = [str(row_index), values[0] + " — " + values[1], values[2]]
        for column_index, value in enumerate(row_values):
            fill_template_cell(table.cell(row_index, column_index), value)


def fill_template_body(doc: Document, content: dict) -> None:
    paragraphs = doc.paragraphs
    content_table_prototype = doc.tables[7]
    toc_lines = (
        ["OVERVIEW", "• Glossary", "• IDTS Solution Context and Scope", "ORGANIZATIONAL STRUCTURE", "BUSINESS PROCESS", "• IDTS-BP-01 Issue and Defect Management", "• Process Flow", "• Process Description", "REPORTS"]
        if content is EN
        else ["TỔNG QUAN (OVERVIEW)", "• Thuật ngữ", "• Bối cảnh giải pháp và phạm vi IDTS", "CƠ CẤU TỔ CHỨC", "QUY TRÌNH NGHIỆP VỤ", "• IDTS-BP-01 Quản lý Issue và Defect", "• Luồng quy trình", "• Mô tả quy trình", "BÁO CÁO (REPORTS)"]
    )
    set_template_paragraph(paragraphs[23], "\n".join(toc_lines))
    set_template_paragraph(paragraphs[25], "OVERVIEW" if content is EN else "TỔNG QUAN (OVERVIEW)", "Heading 1")
    set_template_paragraph(paragraphs[26], "Glossary" if content is EN else "Thuật ngữ", "Heading 2")
    set_template_paragraph(paragraphs[27], "IDTS Solution Context and Scope" if content is EN else "Bối cảnh giải pháp và phạm vi IDTS", "Heading 2")
    objective_label = "Business objectives" if content is EN else "Mục tiêu nghiệp vụ"
    scope_label = "In scope" if content is EN else "Trong phạm vi"
    out_label = "Out of scope" if content is EN else "Ngoài phạm vi"
    set_template_paragraph(paragraphs[28], content["purpose_text"])
    set_template_paragraph(paragraphs[29], objective_label + ":\n" + "\n".join(f"• {item}" for item in content["objectives_items"]) + "\n\n" + scope_label + ":\n" + "\n".join(f"• {item}" for item in content["in_scope_items"]) + "\n\n" + out_label + ":\n" + "\n".join(f"• {item}" for item in content["out_scope_items"]))
    overview_anchor = add_table_series_after(doc, paragraphs[30], overview_table_specs(content), content_table_prototype, 3)
    set_template_paragraph(paragraphs[31], "", "Normal")

    set_template_paragraph(paragraphs[32], "ORGANIZATIONAL STRUCTURE" if content is EN else "CƠ CẤU TỔ CHỨC (ORGANIZATIONAL STRUCTURE)", "Heading 1")
    set_template_paragraph(paragraphs[33], content["organization_text"])
    organization_anchor = add_table_series_after(doc, paragraphs[33], organization_table_specs(content), content_table_prototype, 2)

    set_template_paragraph(paragraphs[34], "BUSINESS PROCESS" if content is EN else "QUY TRÌNH NGHIỆP VỤ (BUSINESS PROCESS)", "Heading 1")
    set_template_paragraph(paragraphs[35], "IDTS-BP-01 Issue and Defect Management" if content is EN else "IDTS-BP-01 Quản lý Issue và Defect", "Heading 2")
    set_template_paragraph(paragraphs[36], "Process Flow" if content is EN else "Luồng quy trình", "Heading 2")
    set_template_paragraph(paragraphs[37], content["process_intro"])
    lifecycle_anchor = add_table_series_after(doc, paragraphs[37], [lifecycle_table_spec(content)], content_table_prototype, 3)
    set_template_paragraph(paragraphs[38], "Process Description" if content is EN else "Mô tả quy trình", "Heading 2")
    paragraphs[38].paragraph_format.page_break_before = True

    set_template_paragraph(paragraphs[39], "REPORTS" if content is EN else "BÁO CÁO (REPORTS)", "Heading 1")
    paragraphs[39].paragraph_format.page_break_before = True
    set_template_paragraph(paragraphs[40], "The following IDTS views support operational monitoring and mentor review." if content is EN else "Các khung nhìn IDTS sau hỗ trợ theo dõi vận hành và việc người hướng dẫn đánh giá.")
    supplemental = (
        "Security, quality and controls:\n" + rows_to_lines(content["quality_rows"]) + "\n\nTraceability and evidence:\n" + rows_to_lines(content["trace_rows"])
        if content is EN
        else "An toàn, chất lượng và kiểm soát:\n" + rows_to_lines(content["quality_rows"]) + "\n\nTruy vết và bằng chứng:\n" + rows_to_lines(content["trace_rows"])
    )
    set_template_paragraph(paragraphs[41], supplemental)
    report_anchor = add_table_series_after(doc, paragraphs[41], report_table_specs(content), content_table_prototype, 2)

    insert_picture_after(doc, overview_anchor, "01-system-context", "Figure 1 — IDTS system context and external boundaries" if content is EN else "Hình 1 — Bối cảnh hệ thống và ranh giới bên ngoài của IDTS")
    anchor = insert_picture_after(doc, lifecycle_anchor, "20-frs-resolve-retest-close-reopen", "Figure 2 — Resolve, retest, close and reopen flow" if content is EN else "Hình 2 — Luồng xác nhận xử lý, kiểm thử lại, đóng và mở lại")
    insert_picture_after(doc, anchor, "07-developer-review", "Figure 3 — Developer review and controlled follow-up" if content is EN else "Hình 3 — Developer đánh giá và xử lý tiếp có kiểm soát")
    if len(doc.tables) != 18:
        raise ValueError(f"Blueprint output contract changed: expected 8 preserved template tables plus 10 approved content tables, found {len(doc.tables)}")


def build(content: dict) -> Path:
    if not TEMPLATE.exists() or TEMPLATE.stat().st_size == 0:
        raise FileNotFoundError(f"Missing official Blueprint template: {TEMPLATE}")
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUT_DIR / content["output"]
    shutil.copy2(TEMPLATE, output)
    doc = Document(output)
    if len(doc.sections) != 3 or len(doc.tables) != 8:
        raise ValueError("Official Blueprint template contract changed: expected 3 sections and 8 body tables")
    enable_field_updates(doc)
    remove_sample_objects(doc)
    fill_cover(doc, content)
    fill_control_tables(doc, content)
    fill_glossary(doc, content)
    fill_process_table(doc, content)
    fill_reports_table(doc, content)
    fill_template_body(doc, content)
    replacements = {
        "FPT Software HCM Co., Ltd.": "IDTS SAP01",
        "Created By Van Bao Chau": "Prepared by DonHV",
        "Created By AAAA": "Prepared by DonHV",
        "Created By ": "Prepared by ",
        "AAAA": "DonHV",
        "Nguyen Hoang Group": "Prepared by DonHV",
    }
    for relationship in doc.part.rels.values():
        if relationship.reltype == RT.HEADER:
            replace_part_text(relationship.target_part, replacements)
        elif relationship.reltype == RT.FOOTER:
            replace_part_text(relationship.target_part, {"Confidental": "Confidential"})
    doc.core_properties.title = f"IDTS SAP490 Blueprint {content['language']} {VERSION}"
    doc.core_properties.subject = "Official-template-filled IDTS SAP490 Blueprint"
    doc.core_properties.author = "IDTS SAP01 Team"
    doc.core_properties.last_modified_by = "IDTS SAP01 Team"
    doc.core_properties.comments = "Template-filled local review candidate; no credentials or private endpoints."
    doc.save(output)
    return output


def main() -> None:
    for content in (EN, VI):
        output = build(content)
        print(output.relative_to(ROOT))


if __name__ == "__main__":
    main()
