"""Validate current Blueprint v0.6 invariants without enforcing obsolete v0.4 layout details."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT / "docs/sap490/templates/Deliverable_template/Blueprint_Template.docx"
OUTPUTS = {
    "en": ROOT / "docs/sap490/generated/Blueprint_IDTS_SAP01_en_v0.6.docx",
}
FORBIDDEN = ("AAAA", "Created By Van Bao Chau", "Nguyen Hoang Group")


def style_signature(document: Document) -> list[tuple[str, str, int]]:
    return [(style.style_id, style.name, int(style.type)) for style in document.styles]


def all_text(document: Document) -> str:
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def validate(language: str, path: Path, template: Document) -> tuple[list[str], int]:
    errors: list[str] = []
    if not path.exists() or path.stat().st_size == 0:
        return [f"{language}: missing or empty output {path}"], 0

    document = Document(path)
    text = all_text(document)
    if len(document.sections) != len(template.sections):
        errors.append(
            f"{language}: section count {len(document.sections)} differs from template {len(template.sections)}"
        )
    if len(document.tables) < len(template.tables):
        errors.append(
            f"{language}: table count {len(document.tables)} is below template core count {len(template.tables)}"
        )
    if style_signature(document) != style_signature(template):
        errors.append(f"{language}: style IDs/names/types differ from the official template")
    if "v0.6" not in path.name or "v0.6" not in text:
        errors.append(f"{language}: filename/content version is not v0.6")
    for placeholder in FORBIDDEN:
        if placeholder in text:
            errors.append(f"{language}: sample/placeholder text remains: {placeholder!r}")

    bp_numbers = sorted({int(value) for value in re.findall(r"BP-(\d{2})", text)})
    expected = list(range(1, 14))
    if bp_numbers != expected:
        errors.append(f"{language}: BP coverage is {bp_numbers}; expected {expected}")
    return errors, len(document.tables)


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    if not TEMPLATE.exists():
        print(f"FAIL: missing official template {TEMPLATE}")
        return 1

    template = Document(TEMPLATE)
    if len(template.sections) != 3 or len(template.tables) != 8:
        print(
            "FAIL: official template baseline changed "
            f"(sections={len(template.sections)}, tables={len(template.tables)})"
        )
        return 1

    failures: list[str] = []
    table_counts: dict[str, int] = {}
    for language, path in OUTPUTS.items():
        errors, table_count = validate(language, path, template)
        failures.extend(errors)
        table_counts[language] = table_count
    if failures:
        print(f"FAIL: Blueprint validation found {len(failures)} issue(s)")
        for failure in failures:
            print(f"- {failure}")
        return 1

    print("PASS: Blueprint v0.6 current-artifact validation")
    print("- official template baseline: 3 sections and 8 core tables")
    print(f"- English output table count: {table_counts['en']}")
    print("- style identity, version, placeholders and BP-01..BP-13 coverage: PASS")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
