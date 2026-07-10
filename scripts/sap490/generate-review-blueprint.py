"""Refresh existing template-derived SAP490 Blueprints without rebuilding layout."""

from pathlib import Path
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "sap490" / "generated"


UPDATES = {
    "en": {
        "source": "Blueprint_IDTS_SAP01_en_v0.1.docx",
        "output": "Blueprint_IDTS_SAP01_en_v0.2.docx",
        "baseline_prefix": "Current baseline:",
        "baseline": "Current review baseline: CAP/Fiori MVP is implemented for structured bug reporting, classification, responsibility-aware assignment, lifecycle actions, comments, draft attachments, audit/history, notifications, and PM monitoring. Optional AI suggestions for similar bugs, classification, handoff, and Smart Assign are human-review only; the real provider remains disabled until approved private configuration and live evidence are available.",
        "cover": "SAP490 Project Blueprint | English Version | v0.2",
        "prepared": "Prepared by: DonHV | Date: 2026-07-10 | Template: Blueprint_Template.docx",
        "history": "SAP490 review refresh: synchronized current CAP/Fiori implementation, attachment/QA evidence, and review-only AI boundary.",
    },
    "vi": {
        "source": "Blueprint_IDTS_SAP01_vi_v0.1.docx",
        "output": "Blueprint_IDTS_SAP01_vi_v0.2.docx",
        "baseline_prefix": "Baseline hiện tại:",
        "baseline": "Baseline review hiện tại: CAP/Fiori MVP đã triển khai bug reporting có cấu trúc, classification, assignment theo responsibility, lifecycle action, comment, draft attachment, audit/history, notification và PM monitoring. AI suggestion tùy chọn cho similar bug, classification, handoff và Smart Assign chỉ để human review; real provider vẫn tắt cho đến khi có private configuration được duyệt và live evidence.",
        "cover": "SAP490 Project Blueprint | Phiên bản tiếng Việt | v0.2",
        "prepared": "Người chuẩn bị: DonHV | Ngày: 2026-07-10 | Template: Blueprint_Template.docx",
        "history": "Refresh SAP490 review: đồng bộ CAP/Fiori implementation hiện tại, evidence attachment/QA và AI chỉ để review.",
    },
}


def set_text(paragraph, text):
    if not paragraph.runs:
        paragraph.add_run()
    first = paragraph.runs[0]
    for run in paragraph.runs:
        run.text = ""
    first.text = text


def set_cell(cell, text):
    set_text(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        set_text(paragraph, "")


def build(language):
    values = UPDATES[language]
    source = OUT / values["source"]
    output = OUT / values["output"]
    shutil.copy2(source, output)
    doc = Document(output)

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(values["baseline_prefix"]):
            set_text(paragraph, values["baseline"])
            break
    else:
        raise ValueError(f"Baseline paragraph not found in {source}")

    set_cell(doc.tables[0].rows[1].cells[0], values["cover"])
    set_cell(doc.tables[0].rows[3].cells[0], values["prepared"])
    # Reuse an existing history row to preserve original table layout.
    history_row = doc.tables[1].rows[3].cells
    history_values = ["2026-07-10", "SAP490 review refresh", values["history"], "DonHV / Codex", "C", "v0.2"]
    for cell, value in zip(history_row, history_values):
        set_cell(cell, value)

    doc.core_properties.title = f"IDTS SAP490 Blueprint {language.upper()} v0.2"
    doc.core_properties.subject = "Current implementation review; no secrets"
    doc.core_properties.author = "IDTS SAP01 Team"
    doc.save(output)
    return output


def main():
    for language in ("en", "vi"):
        print(build(language).relative_to(ROOT))


if __name__ == "__main__":
    main()
