"""Build review-ready BRD/SRS/FRS DOCX files from canonical Markdown and PNG figures."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS = [
    ("docs/ba/brd/brd.en.md", "docs/ba/brd/brd.en.docx"),
    ("docs/ba/brd/brd.vi.md", "docs/ba/brd/brd.vi.docx"),
    ("docs/ba/srs/srs.en.md", "docs/ba/srs/srs.en.docx"),
    ("docs/ba/srs/srs.vi.md", "docs/ba/srs/srs.vi.docx"),
    ("docs/ba/frs/frs.en.md", "docs/ba/frs/frs.en.docx"),
    ("docs/ba/frs/frs.vi.md", "docs/ba/frs/frs.vi.docx"),
]
IMAGE = re.compile(r"^!\[(?P<alt>[^]]*)\]\((?P<path>[^)]+)\)$")
HEADING = re.compile(r"^(?P<hashes>#{1,6})\s+(?P<text>.+)$")


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(clean_inline(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    cell.vertical_alignment = 1


def clean_inline(value: str) -> str:
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = value.replace("**", "").replace("*", "")
    return value.replace("\\|", "|").strip()


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)
    vertical_alignment = tc_pr.find(qn("w:vAlign"))
    if vertical_alignment is None:
        tc_pr.append(shading)
    else:
        tc_pr.insert(list(tc_pr).index(vertical_alignment), shading)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def add_table(document: Document, rows: list[str]) -> None:
    data = [[clean_inline(part) for part in row.strip().strip("|").split("|")] for row in rows]
    if len(data) < 2:
        return
    headers = data[0]
    body = [row for row in data[2:] if any(cell.strip() for cell in row)]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        shade_cell(table.rows[0].cells[index], "D9E2F3")
    for row in body:
        cells = table.add_row().cells
        for index in range(len(headers)):
            set_cell_text(cells[index], row[index] if index < len(row) else "")
    # Do not add a blank paragraph after each table: OfficeCLI correctly flags
    # those empty paragraphs as presentation issues. Word separates adjacent
    # blocks without it, and the next meaningful Markdown block supplies space.


def add_figure(document: Document, markdown_path: Path, alt: str, source: str) -> None:
    asset_name = Path(source).stem
    png = ROOT / "docs/diagrams/rendered/png" / f"{asset_name}.png"
    if not png.exists():
        raise RuntimeError(f"Missing PNG for figure: {png}")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(png), width=Inches(6.3))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(f"Figure. {alt}")
    run.italic = True
    run.font.size = Pt(9)


def apply_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    for level, size in [(1, 18), (2, 14), (3, 12), (4, 11)]:
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)


def remove_empty_body_paragraphs(document: Document) -> None:
    """Remove structural-only paragraphs that OfficeCLI flags as visual noise.

    Paragraphs with a drawing or page-break still have a run and are retained;
    only truly empty body paragraphs are removed.
    """
    for paragraph in list(document.paragraphs):
        has_run = bool(paragraph._p.findall(qn("w:r")))
        if not paragraph.text and not has_run:
            paragraph._element.getparent().remove(paragraph._element)


def build(markdown_path: Path, output_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    apply_base_styles(document)

    title = lines[0].lstrip("# ").strip()
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    for line in lines[2:8]:
        if line.strip():
            paragraph = document.add_paragraph(clean_inline(line))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()

    index = 0
    while index < len(lines):
        line = lines[index]
        if index < 8:
            index += 1
            continue
        image = IMAGE.match(line)
        if image:
            add_figure(document, markdown_path, image.group("alt"), image.group("path"))
            index += 1
            continue
        heading = HEADING.match(line)
        if heading:
            level = min(len(heading.group("hashes")), 4)
            document.add_heading(clean_inline(heading.group("text")), level=level)
            index += 1
            continue
        if line.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, table_lines)
            continue
        if not line.strip() or line.startswith("*Figure "):
            index += 1
            continue
        if line.startswith(("- ", "1. ", "2. ", "3. ", "4. ", "5. ", "6. ", "7. ", "8. ", "9. ")):
            paragraph = document.add_paragraph(style="List Bullet" if line.startswith("- ") else "List Number")
            paragraph.add_run(clean_inline(re.sub(r"^(?:- |\d+\. )", "", line)))
        else:
            document.add_paragraph(clean_inline(line))
        index += 1

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("IDTS SAP490 Review | Page ")
    add_page_field(footer)
    remove_empty_body_paragraphs(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    for source, target in DOCS:
        build(ROOT / source, ROOT / target)
        print(f"Built {target}")


if __name__ == "__main__":
    main()
