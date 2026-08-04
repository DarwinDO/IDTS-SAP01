"""Red/green quality contract for the mentor-facing SAP490 specifications.

This test intentionally checks content truth in addition to Office file
validity. It prevents a schema-valid English submission from being accepted
when its runtime trace, test status, or version history is stale.
"""

from __future__ import annotations

from pathlib import Path
from functools import lru_cache
import re
import sys

from docx import Document
from openpyxl import load_workbook
from specification_catalog import FUNCTIONS, LIFECYCLE_ACTIONS, MESSAGES, TECH_REQUIREMENTS


ROOT = Path(__file__).resolve().parents[2]
GENERATED = ROOT / "docs" / "sap490" / "generated"
TECHNICAL_TEMPLATE = ROOT / "docs" / "sap490" / "templates" / "Deliverable_template" / "Technical_Specification.xlsx"
TECHNICAL_MESSAGE_SOURCE = ROOT / "docs" / "pm" / "evidence" / "idts-109" / "technical-spec" / "message-catalog.md"

FILES = {
    "functional_en": GENERATED / "Functional_Specification_IDTS_SAP01_en_v0.7.xlsx",
    "technical_en": GENERATED / "Technical_Specification_IDTS_SAP01_en_v0.8.xlsx",
    "config_en": GENERATED / "Configuration_Note_IDTS_SAP01_en_v0.5.xlsx",
    "blueprint_en": GENERATED / "Blueprint_IDTS_SAP01_en_v0.6.docx",
}

FORBIDDEN_TRACES = {
    "srv/auth-service.js",
    "srv/bug-service/read-models.js",
    "LoginPage.js",
    "ProfileMenu.js",
    "Dashboard.view.xml",
    "Dashboard.controller.js",
    "*Review.js",
}

REQUIRED_TRACES = {
    "srv/auth.js",
    "srv/bug-service/monitoring.js",
    "login-page.js",
    "LoginController.js",
    "ProfileShell.js",
    "dashboard.html",
    "dashboard-page.js",
    "ClassificationReview.js",
    "DuplicateReview.js",
    "HandoffSummaryReview.js",
    "SmartAssignDeveloper.js",
    "BugCollaboration.js",
}

REQUIRED_TEST_TRUTH = {
    "38 accepted candidates",
    "2 held",
    "135 mapping-only",
    "13 blocked",
    "22 MEETS",
    "12 DOES_NOT_MEET",
    "23 BLOCKED",
    "not final acceptance claims",
}


@lru_cache(maxsize=None)
def workbook_text(path: Path) -> str:
    workbook = load_workbook(path, data_only=False)
    return "\n".join(
        str(cell.value)
        for sheet in workbook.worksheets
        for row in sheet.iter_rows()
        for cell in row
        if cell.value not in (None, "")
    )


@lru_cache(maxsize=None)
def document_text(path: Path) -> str:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    values.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(values)


def visible_cells(path: Path) -> list:
    workbook = load_workbook(path, data_only=False)
    return [
        cell
        for sheet in workbook.worksheets
        if sheet.sheet_state == "visible"
        for row in sheet.iter_rows()
        for cell in row
        if cell.value not in (None, "")
    ]


def validate_technical_template_contract(path: Path, language: str) -> list[str]:
    """Protect the official inner layouts, not only tab names and outer styling."""
    failures = []
    template = load_workbook(TECHNICAL_TEMPLATE, data_only=False)
    workbook = load_workbook(path, data_only=False)
    expected_sheets = template.sheetnames
    if workbook.sheetnames != expected_sheets:
        failures.append(f"{path.name}: sheet order differs from official template")
    for name in expected_sheets:
        if workbook[name].sheet_state != template[name].sheet_state:
            failures.append(f"{path.name}/{name}: visibility differs from official template")

    layout = workbook["Screen Layout"]
    template_layout = template["Screen Layout"]
    required_layout_merges = {"B5:BG5", "B6:BG6", "B7:BG7", "C8:BG8"}
    if not required_layout_merges.issubset({str(item) for item in layout.merged_cells.ranges}):
        failures.append(f"{path.name}: Screen Layout official merged blocks were not preserved")
    if len(layout._images) < 2:
        failures.append(f"{path.name}: Screen Layout must contain two implemented-screen images")
    layout_text = "\n".join(str(cell.value) for row in layout.iter_rows() for cell in row if cell.value)
    for forbidden in ("Route / entry", "Controller / extension", "OData binding"):
        if forbidden in layout_text:
            failures.append(f"{path.name}: Screen Layout still contains custom source-map table header {forbidden!r}")

    definition = workbook["Screen Definition"]
    if str(definition["B9"].value or "").startswith("Screen ID"):
        failures.append(f"{path.name}: Screen Definition custom table replaced the official field grid")
    for coordinate in ("B10", "C11", "I11", "M11", "O11", "R11", "AW10"):
        if definition[coordinate].value in (None, ""):
            failures.append(f"{path.name}: Screen Definition official header {coordinate} is empty")

    messages = workbook["Message Definition"]
    expected_headers = {
        "Message ID", "User-facing text / safe summary", "Exact trigger / source",
        "HTTP / status", "Target", "Role / context", "Rollback behavior",
        "Sanitized logging", "Frontend handling", "Evidence",
    }
    actual_headers = {
        str(cell.value) for cell in messages[5] if cell.value not in (None, "")
    }
    if actual_headers != expected_headers:
        failures.append(
            f"{path.name}: Message Definition formal ten-column catalog is incomplete: "
            f"{sorted(expected_headers ^ actual_headers)}"
        )
    message_ids = [
        str(messages.cell(row, 2).value)
        for row in range(6, messages.max_row + 1)
        if str(messages.cell(row, 2).value or "").startswith("MSG-")
    ]
    source_message_ids = {
        match.group(1)
        for match in re.finditer(r"^\|\s*(MSG-[A-Z0-9-]+)\s*\|", TECHNICAL_MESSAGE_SOURCE.read_text(encoding="utf-8"), re.MULTILINE)
    }
    if set(message_ids) != source_message_ids or len(message_ids) != len(set(message_ids)):
        failures.append(
            f"{path.name}: Message Definition differs from its current source catalog; "
            f"workbook={len(message_ids)} rows/{len(set(message_ids))} unique, "
            f"source={len(source_message_ids)} unique, delta={sorted(set(message_ids) ^ source_message_ids)}"
        )
    for row in range(6, messages.max_row + 1):
        if not str(messages.cell(row, 2).value or "").startswith("MSG-"):
            continue
        values = [messages.cell(row, col).value for col in (2, 7, 17, 26, 30, 34, 38, 42, 47, 53)]
        if any(value in (None, "") for value in values):
            failures.append(f"{path.name}: incomplete Message Definition record at row {row}")

    implementation_text = "\n".join(
        str(cell.value) for row in workbook["Technical Implementation"].iter_rows()
        for cell in row if cell.value not in (None, "")
    )
    required_flows = {"FLOW-COMMENT-CREATE", "TI-COLLAB-02"}
    required_flows.update(f"FLOW-{action.upper()}" for action, *_ in LIFECYCLE_ACTIONS)
    for flow in sorted(required_flows):
        if flow not in implementation_text:
            failures.append(f"{path.name}: missing exact implementation flow {flow}")
    for retired_flow in ("FLOW-ATTACH-QUEUE", "FLOW-ATTACH-UPLOAD"):
        if retired_flow in implementation_text:
            failures.append(f"{path.name}: retired attachment flow remains: {retired_flow}")
    for forbidden in ("Add comment or attachment", "POST/PUT/DELETE /odata/v4/bug child entity"):
        if forbidden in implementation_text:
            failures.append(f"{path.name}: generic or command-only implementation evidence remains: {forbidden}")
    full_text = workbook_text(path)
    for forbidden in ("WS92400001", "Credit Memo Request", "NamNH", "HuyNB", "ThaoML9"):
        if forbidden in full_text:
            failures.append(f"{path.name}: official-template sample residue remains: {forbidden}")

    if language == "vi":
        for forbidden in ("Route / entry", "Page / view", "Role-scoped", "Signed-in user", "Bug participants", "Anonymous"):
            if forbidden in workbook_text(path):
                failures.append(f"{path.name}: visible English reader-facing label remains: {forbidden}")
    return failures


def validate_catalog_sources() -> list[str]:
    failures = []
    references = [item["source"] for item in FUNCTIONS]
    references.extend(item["source"] for item in MESSAGES)
    references.extend(item[5] for item in TECH_REQUIREMENTS)
    for reference in references:
        for segment in re.split(r"\s*[;→]\s*", reference):
            if not segment.startswith(("app/", "srv/", "db/")):
                continue
            path_text, _, symbol = segment.partition("::")
            path = ROOT / path_text.rstrip("/.,")
            if not path.exists():
                failures.append(f"catalog source path does not exist: {path_text}")
                continue
            if symbol and path.is_file() and symbol not in path.read_text(encoding="utf-8"):
                failures.append(f"catalog source symbol does not exist: {reference}")
    return failures


def main() -> int:
    failures: list[str] = validate_catalog_sources()
    for label, path in FILES.items():
        if not path.exists():
            failures.append(f"{label}: missing expected versioned artifact {path.name}")

    if failures:
        print("\n".join(f"FAIL: {item}" for item in failures))
        return 1

    failures.extend(validate_technical_template_contract(FILES["technical_en"], "en"))
    functional_text = workbook_text(FILES["functional_en"])
    technical_text = workbook_text(FILES["technical_en"])
    config_text = workbook_text(FILES["config_en"])
    blueprint_text = document_text(FILES["blueprint_en"])
    all_text = "\n".join((functional_text, technical_text, config_text, blueprint_text))

    for trace in sorted(FORBIDDEN_TRACES):
        if trace in all_text:
            failures.append(f"forbidden stale runtime trace remains: {trace}")
    for trace in sorted(REQUIRED_TRACES):
        if trace not in all_text:
            failures.append(f"missing exact runtime trace: {trace}")
    for truth in sorted(REQUIRED_TEST_TRUTH):
        if truth not in technical_text:
            failures.append(f"Technical Specification missing current reviewed test truth: {truth}")

    for requirement in ("SRS-FR-ASG", "SRS-FR-MON"):
        if requirement not in technical_text:
            failures.append(f"Technical Specification missing requirement group {requirement}")

    if "token returned once" not in all_text.casefold() and "raw token is returned once" not in all_text.casefold():
        failures.append("AuthSessions wording does not explain one-time raw token return")
    if "tokenHash" not in all_text:
        failures.append("AuthSessions wording does not name tokenHash")
    for attachment_trace in ("@cap-js/attachments", "BugAttachments"):
        if attachment_trace not in technical_text:
            failures.append(f"Technical Specification missing standard attachment trace: {attachment_trace}")
    for retired_trace in ("pendingCreateAttachmentsByBugId", "queuePendingCreateAttachments", "uploadFilesToSavedBug"):
        if retired_trace in technical_text:
            failures.append(f"Technical Specification contains retired custom attachment trace: {retired_trace}")

    generic_409_patterns = (
        "Stale/repeated transition",
        "The bug state changed. Reload and try again. [HTTP 409]",
        "Trạng thái Bug đã thay đổi. Hãy tải lại và thử lại. [HTTP 409]",
    )
    if any(pattern in all_text for pattern in generic_409_patterns):
        failures.append("HTTP 409 is still described as a generic lifecycle response")

    for path in (FILES["technical_en"],):
        workbook = load_workbook(path, data_only=False)
        implementation = workbook["Technical Implementation"]
        bad_fonts = [
            cell.coordinate
            for row in implementation.iter_rows()
            for cell in row
            if cell.value not in (None, "")
            and (cell.font.name or "").casefold() == "calibri"
        ]
        if bad_fonts:
            failures.append(f"{path.name}: non-template Technical Implementation font at {bad_fonts[:8]}")

    for path in FILES.values():
        text = workbook_text(path) if path.suffix == ".xlsx" else document_text(path)
        if "###" in text:
            failures.append(f"{path.name}: contains Excel overflow marker ###")
        if "Created bDonHV" in text or "IDTSIDTS-TECH" in text:
            failures.append(f"{path.name}: contains concatenated metadata")

    for path in (FILES["functional_en"],):
        workbook = load_workbook(path, data_only=False)
        definition = workbook["Screen Definition"]
        populated = sum(
            1
            for row in definition.iter_rows(min_row=12)
            if any(cell.value not in (None, "") for cell in row)
        )
        if populated < 18:
            failures.append(f"{path.name}: Screen Definition has only {populated} populated rows; expected >=18")

    formal_paths = (FILES["functional_en"], FILES["technical_en"])
    for path in formal_paths:
        workbook = load_workbook(path, data_only=False)
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    value = str(cell.value or "")
                    if " | " in value:
                        failures.append(f"{path.name}/{sheet.title}!{cell.coordinate}: raw pipe-delimited record")
                    is_function_data = sheet.title != "Function Overview" or cell.row >= 15
                    is_requirement_data = sheet.title != "Function Overview" or cell.row >= 15
                    if is_function_data and len(re.findall(r"\bFN-[A-Z]+-\d+\b", value)) > 1:
                        failures.append(f"{path.name}/{sheet.title}!{cell.coordinate}: multiple Function IDs in one cell")
                    if is_requirement_data and len(re.findall(r"\bSRS-FR-[A-Z]+\b", value)) > 1:
                        failures.append(f"{path.name}/{sheet.title}!{cell.coordinate}: multiple Requirement IDs in one cell")

    for path in (FILES["technical_en"],):
        workbook = load_workbook(path, data_only=False)
        layout_text = "\n".join(
            str(cell.value) for row in workbook["Screen Layout"].iter_rows()
            for cell in row if cell.value not in (None, "")
        )
        if "→" in layout_text or "Route/Page/Extension map" in layout_text:
            failures.append(f"{path.name}: Screen Layout still contains prose trace instead of screen records")
        if "IDTS-TECH-" in workbook_text(path):
            failures.append(f"{path.name}: independent technical message catalog remains")

    # Functional Specification is no longer a required mentor artifact.  The
    # Technical Specification owns the exhaustive current message catalog and
    # is validated above against its 145-row source inventory instead of a
    # frozen legacy Functional workbook.

    if failures:
        print("\n".join(f"FAIL: {item}" for item in failures))
        return 1
    print("PASS: SAP490 specification quality contract")
    return 0


if __name__ == "__main__":
    sys.exit(main())
