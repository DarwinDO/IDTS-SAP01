"""Validate the mentor-facing SAP490 specification pack.

The validator is intentionally strict about official-template fidelity.  It
checks the eight generated EN/VI artifacts without modifying them.
"""

from __future__ import annotations

from collections import Counter
from pathlib import Path
from zipfile import ZipFile
import re
import sys

from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[2]
TEMPLATE_DIR = ROOT / "docs" / "sap490" / "templates" / "Deliverable_template"
GENERATED_DIR = ROOT / "docs" / "sap490" / "generated"

XLSX_CONTRACTS = {
    "functional": {
        "template": TEMPLATE_DIR / "Functional_Specification.xlsx",
        "outputs": [
            GENERATED_DIR / "Functional_Specification_IDTS_SAP01_en_v0.7.xlsx",
            GENERATED_DIR / "Functional_Specification_IDTS_SAP01_vi_v0.7.xlsx",
        ],
        "required_sheets": [
            "Cover",
            "Histories",
            "Function Overview",
            "Process Flow",
            "Screen Layout",
            "Screen Definition",
            "Smart Form Structure",
            "Message Definition",
            "Processing Description",
        ],
    },
    "technical": {
        "template": TEMPLATE_DIR / "Technical_Specification.xlsx",
        "outputs": [
            GENERATED_DIR / "Technical_Specification_IDTS_SAP01_en_v0.6.xlsx",
            GENERATED_DIR / "Technical_Specification_IDTS_SAP01_vi_v0.6.xlsx",
        ],
        "required_sheets": [
            "Cover",
            "Histories",
            "Introduction",
            "Scope",
            "Assumptions",
            "Functional Requirements",
            "Technical Design",
            "Development Standards",
            "Screen Layout",
            "Screen Definition",
            "Message Definition",
            "Technical Implementation",
        ],
    },
    "configuration": {
        "template": TEMPLATE_DIR / "Configuration_Note.xlsx",
        "outputs": [
            GENERATED_DIR / "Configuration_Note_IDTS_SAP01_en_v0.5.xlsx",
            GENERATED_DIR / "Configuration_Note_IDTS_SAP01_vi_v0.5.xlsx",
        ],
        "required_sheets": ["Cover", "Record of change", "Checklist", "4", "5"],
    },
}

BLUEPRINT_TEMPLATE = TEMPLATE_DIR / "Blueprint_Template.docx"
BLUEPRINT_OUTPUTS = [
    GENERATED_DIR / "Blueprint_IDTS_SAP01_en_v0.6.docx",
    GENERATED_DIR / "Blueprint_IDTS_SAP01_vi_v0.6.docx",
]

FORBIDDEN_RESIDUE = [
    "VBAK",
    "VBAP",
    "KNA1",
    "KNVV",
    "Sales Data",
    "WS92400001",
    "Credit Memo Request",
    "G04 - Workflow",
    "screen 9000",
]

REQUIRED_RUNTIME_TERMS = [
    "/odata/v4/auth/",
    "/odata/v4/bug/",
    "acceptAiSuggestion",
    "rejectAiSuggestion",
    "ignoreAiSuggestion",
    "applyClassificationSuggestion",
    "confirmDuplicateSuggestion",
    "readAiOperationalMetrics",
    "operationStatus",
    "latencyMs",
]

# Official templates contain example data blocks whose merged-cell layout is
# intentionally replaced by record-level tables.  Cover, metadata, sheet order,
# visibility, page setup and columns remain immutable; only these data regions
# may change their row merges and cell styles.
MUTABLE_DATA_REGIONS = {
    "functional": {
        "Function Overview": (15, 90), "Process Flow": (6, 60),
        "Screen Layout": (27, 90), "Screen Definition": (9, 74),
        "Smart Form Structure": (44, 57), "Message Definition": (5, 45),
        "Processing Description": (6, 90),
    },
    "technical": {
        "Introduction": (15, 30), "Scope": (6, 40), "Assumptions": (6, 50),
        "Functional Requirements": (5, 75), "Technical Design": (5, 129),
        "Development Standards": (5, 95), "Screen Layout": (5, 40),
        "Screen Definition": (9, 50), "Message Definition": (5, 40),
        "Technical Implementation": (5, 40),
    },
}


def row_in_mutable_region(kind, sheet, row):
    region = MUTABLE_DATA_REGIONS.get(kind, {}).get(sheet)
    return bool(region and region[0] <= row <= region[1])


def color_signature(color):
    if color is None:
        return None
    return (
        color.type,
        color.rgb,
        color.indexed,
        color.theme,
        round(color.tint or 0, 6),
    )


def visual_style_signature(cell):
    """Compare visible style, ignoring package-only Excel defaults."""

    font = cell.font
    fill = cell.fill
    border = cell.border
    alignment = cell.alignment
    return (
        font.name,
        font.sz,
        bool(font.bold),
        bool(font.italic),
        font.underline or None,
        bool(font.strike),
        color_signature(font.color),
        fill.fill_type,
        color_signature(fill.fgColor),
        border.left.style,
        border.right.style,
        border.top.style,
        border.bottom.style,
        alignment.horizontal or "general",
        alignment.vertical or "bottom",
        bool(alignment.wrap_text),
        alignment.text_rotation or 0,
        bool(alignment.shrink_to_fit),
    )


def visible_style_matches(template_cell, output_cell):
    """Allow only the documented white-theme-to-black readability correction."""
    expected = visual_style_signature(template_cell)
    actual = visual_style_signature(output_cell)
    if expected == actual:
        return True
    template_color = template_cell.font.color
    output_color = output_cell.font.color
    readability_override = (
        output_cell.value not in (None, "")
        and template_color is not None
        and template_color.type == "theme"
        and template_color.theme == 1
        and output_color is not None
        and output_color.type == "rgb"
        and output_color.rgb in {"FF000000", "00000000"}
    )
    return readability_override and expected[:6] == actual[:6] and expected[7:] == actual[7:]


def page_signature(worksheet):
    margins = worksheet.page_margins
    setup = worksheet.page_setup
    return (
        setup.orientation,
        setup.paperSize,
        round(margins.left or 0, 2),
        round(margins.right or 0, 2),
        round(margins.top or 0, 2),
        round(margins.bottom or 0, 2),
        round(margins.header or 0, 2),
        round(margins.footer or 0, 2),
        worksheet.print_title_rows,
        worksheet.print_title_cols,
        worksheet.freeze_panes,
        True if worksheet.sheet_view.showGridLines is None else worksheet.sheet_view.showGridLines,
        worksheet.oddHeader.left.text,
        worksheet.oddHeader.center.text,
        worksheet.oddHeader.right.text,
        worksheet.oddFooter.left.text,
        worksheet.oddFooter.center.text,
        worksheet.oddFooter.right.text,
    )


def page_signature_matches(actual, expected):
    """Permit only the approved repair of the template's incomplete footer."""
    if actual == expected:
        return True
    return (
        actual[:-1] == expected[:-1]
        and expected[-1] in {None, "&P / "}
        and actual[-1] == "&P / &N"
    )


def column_width_signature(worksheet):
    return {
        key: (
            round(dimension.width or 0, 2),
            dimension.hidden,
            dimension.outlineLevel,
        )
        for key, dimension in worksheet.column_dimensions.items()
    }


def column_width_signature_matches(actual, expected):
    if set(actual) != set(expected):
        return False
    return all(
        abs(actual[key][0] - expected[key][0]) <= 0.02
        and actual[key][1:] == expected[key][1:]
        for key in expected
    )


def workbook_text(workbook) -> str:
    return "\n".join(
        str(cell.value)
        for worksheet in workbook.worksheets
        for row in worksheet.iter_rows()
        for cell in row
        if cell.value not in (None, "")
    )


def duplicate_drawing_ids(path: Path) -> list[str]:
    duplicates: list[str] = []
    with ZipFile(path) as archive:
        for name in archive.namelist():
            if not name.startswith("xl/drawings/drawing") or not name.endswith(".xml"):
                continue
            xml = archive.read(name).decode("utf-8", errors="replace")
            ids = re.findall(r"<xdr:cNvPr[^>]*\bid=\"([^\"]+)\"", xml)
            repeated = sorted(key for key, count in Counter(ids).items() if count > 1)
            if repeated:
                duplicates.append(f"{name}: {', '.join(repeated)}")
    return duplicates


def validate_workbooks() -> list[str]:
    failures: list[str] = []
    for kind, contract in XLSX_CONTRACTS.items():
        template = load_workbook(contract["template"], data_only=False)
        expected_sheets = contract["required_sheets"]
        expected_states = {ws.title: ws.sheet_state for ws in template.worksheets}
        expected_merges = {
            ws.title: {str(item) for item in ws.merged_cells.ranges}
            for ws in template.worksheets
        }
        expected_pages = {
            ws.title: page_signature(ws)
            for ws in template.worksheets
        }
        expected_columns = {
            ws.title: column_width_signature(ws)
            for ws in template.worksheets
        }

        for output in contract["outputs"]:
            if not output.exists():
                failures.append(f"{output.name}: output is missing")
                continue
            workbook = load_workbook(output, data_only=False)
            if workbook.sheetnames != expected_sheets:
                failures.append(
                    f"{output.name}: sheet order {workbook.sheetnames!r} != {expected_sheets!r}"
                )
                continue
            actual_states = {ws.title: ws.sheet_state for ws in workbook.worksheets}
            if actual_states != expected_states:
                failures.append(
                    f"{output.name}: sheet visibility changed: {actual_states!r}"
                )
            for worksheet in workbook.worksheets:
                if not page_signature_matches(page_signature(worksheet), expected_pages[worksheet.title]):
                    failures.append(
                        f"{output.name}/{worksheet.title}: official page setup changed"
                    )
                if not column_width_signature_matches(column_width_signature(worksheet), expected_columns[worksheet.title]):
                    failures.append(
                        f"{output.name}/{worksheet.title}: official column-width contract changed"
                    )
                actual_merges = {str(item) for item in worksheet.merged_cells.ranges}
                protected_merges = {
                    merge for merge in expected_merges[worksheet.title]
                    if not row_in_mutable_region(kind, worksheet.title, template[worksheet.title][merge.split(":")[0]].row)
                }
                if not protected_merges.issubset(actual_merges):
                    missing = sorted(protected_merges - actual_merges)
                    failures.append(
                        f"{output.name}/{worksheet.title}: missing template merges {missing[:5]!r}"
                    )

                template_sheet = template[worksheet.title]
                style_mismatches = []
                for row in template_sheet.iter_rows():
                    for template_cell in row:
                        if template_cell.value in (None, "") or not template_cell.has_style:
                            continue
                        if row_in_mutable_region(kind, worksheet.title, template_cell.row):
                            continue
                        output_cell = worksheet[template_cell.coordinate]
                        if not visible_style_matches(template_cell, output_cell):
                            style_mismatches.append(template_cell.coordinate)
                if style_mismatches:
                    failures.append(
                        f"{output.name}/{worksheet.title}: visible template style changed at "
                        f"{style_mismatches[:8]!r}"
                    )

            text = workbook_text(workbook)
            for residue in FORBIDDEN_RESIDUE:
                if residue.lower() in text.lower():
                    failures.append(f"{output.name}: forbidden template residue {residue!r}")
            if "#REF!" in text:
                failures.append(f"{output.name}: contains #REF!")
            for defined_name in workbook.defined_names.values():
                if "#REF!" in str(defined_name.attr_text):
                    failures.append(
                        f"{output.name}: broken defined name {defined_name.name!r}"
                    )
            if kind in {"functional", "technical"}:
                for term in REQUIRED_RUNTIME_TERMS:
                    if term not in text:
                        failures.append(f"{output.name}: missing runtime trace term {term!r}")
            if kind == "configuration":
                for hidden_sheet in ("4", "5"):
                    if workbook[hidden_sheet].sheet_state != "hidden":
                        failures.append(
                            f"{output.name}/{hidden_sheet}: template hidden state was not preserved"
                        )
                failures.extend(
                    f"{output.name}: duplicate drawing object id {detail}"
                    for detail in duplicate_drawing_ids(output)
                )

            for worksheet in workbook.worksheets:
                for row in worksheet.iter_rows():
                    for cell in row:
                        if isinstance(cell.value, str) and len(cell.value) > 900:
                            failures.append(
                                f"{output.name}/{worksheet.title}!{cell.coordinate}: "
                                f"overlong cell ({len(cell.value)} characters)"
                            )
    return failures


def doc_text(document: Document) -> str:
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return (
        "\n".join(paragraphs + cells)
        .replace("\u200b", "")
        .replace("\u2060", "")
        .replace("\ufeff", "")
    )


def docx_document_xml(path: Path) -> str:
    with ZipFile(path) as archive:
        return archive.read("word/document.xml").decode("utf-8", errors="replace")


def validate_blueprints() -> list[str]:
    failures: list[str] = []
    template = Document(BLUEPRINT_TEMPLATE)
    template_sections = len(template.sections)
    template_styles = len(template.styles)
    template_section_signatures = [
        (
            section.page_width,
            section.page_height,
            section.orientation,
            section.left_margin,
            section.right_margin,
            section.top_margin,
            section.bottom_margin,
            section.header_distance,
            section.footer_distance,
            section.different_first_page_header_footer,
        )
        for section in template.sections
    ]
    template_table_signatures = Counter(
        (
            table.style.name if table.style else None,
            len(table.columns),
        )
        for table in template.tables
    )
    for output in BLUEPRINT_OUTPUTS:
        if not output.exists():
            failures.append(f"{output.name}: output is missing")
            continue
        document = Document(output)
        if len(document.sections) != template_sections:
            failures.append(
                f"{output.name}: sections={len(document.sections)}, expected {template_sections}"
            )
        if len(document.styles) != template_styles:
            failures.append(
                f"{output.name}: styles={len(document.styles)}, expected {template_styles}"
            )
        if len(document.tables) < len(template.tables):
            failures.append(
                f"{output.name}: lost official core tables "
                f"({len(document.tables)} < {len(template.tables)})"
            )
        section_signatures = [
            (
                section.page_width,
                section.page_height,
                section.orientation,
                section.left_margin,
                section.right_margin,
                section.top_margin,
                section.bottom_margin,
                section.header_distance,
                section.footer_distance,
                section.different_first_page_header_footer,
            )
            for section in document.sections
        ]
        if section_signatures != template_section_signatures:
            failures.append(f"{output.name}: official section/page setup changed")
        table_signatures = Counter(
            (
                table.style.name if table.style else None,
                len(table.columns),
            )
            for table in document.tables
        )
        if any(
            table_signatures[signature] < required_count
            for signature, required_count in template_table_signatures.items()
        ):
            failures.append(f"{output.name}: official core-table structure/style changed")
        text = doc_text(document)
        for term in REQUIRED_RUNTIME_TERMS:
            if term not in text:
                failures.append(f"{output.name}: missing runtime trace term {term!r}")
        if "PENDING-only" in text or "remain PENDING" in text:
            failures.append(f"{output.name}: stale PENDING-only AI wording")
        if "Table of Contents" in text or "Mục lục" in text:
            xml = docx_document_xml(output)
            if r"TOC \o" not in xml:
                failures.append(f"{output.name}: TOC field is missing")
    return failures


def main() -> int:
    failures = validate_workbooks() + validate_blueprints()
    if failures:
        print("SAP490 specification validation FAILED")
        for failure in failures:
            print(f"- {failure}")
        return 1
    print("SAP490 specification validation PASSED")
    print("- Functional Specification: 9/9 tabs preserved")
    print("- Technical Specification: 12/12 tabs preserved")
    print("- Configuration Note: 5/5 tabs preserved")
    print("- Blueprint: official section/style/core-table contract preserved")
    return 0


if __name__ == "__main__":
    sys.exit(main())
