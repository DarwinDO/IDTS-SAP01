"""Fill the official FHU Final Project Report template for mentor review.

The output remains a review draft: mentor approval, signatures, and human UAT
results are intentionally left pending. The template, styles, headings, headers,
footers, sections, and core tables are preserved.
"""

from pathlib import Path
import shutil

from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT / "docs/sap490/templates/Deliverable_template/Final Project Report_FHU.docx"
OUTPUT = ROOT / "docs/sap490/generated/Final_Project_Report_IDTS_SAP01_mentor_review_draft_v0.1.docx"


CONTENT = {
    5: "- Ho Chi Minh City, July 2026 -",
    10: "The IDTS SAP01 team thanks the mentor and reviewers for their guidance. This mentor-review draft reports only evidence that can be traced to the repository, Jira, Shared QA, and the SAP490 test pack. Mentor approval, signatures, and human UAT acceptance remain pending.",
    13: "The table below defines the principal terms used by this report. Product codes and technical names are retained where changing them would reduce traceability.",
    16: "This section consolidates the current project baseline from the approved BRD, SRS, FRS, Blueprint, runtime repository, and Shared QA evidence.",
    19: "Project: Issue and Defect Tracking System in SAP (IDTS-SAP01). The solution is an educational SAP CAP Node.js and SAP Fiori Elements application for reporting, classifying, assigning, tracking, notifying, and auditing software defects.",
    21: "Team ownership: DonHV leads BA/PM consolidation, backend integration, deployment, and documentation; DatDT leads Fiori/UI5 work; SangVN supports Object Page and collaboration UX; NhanT leads QA and evidence review. Mentor/supervisor approval remains pending.",
    23: "Software-testing teams need one consistent record of defect facts, ownership, lifecycle state, supporting evidence, comments, notifications, and audit history. Fragmented communication makes assignment, follow-up, and verification difficult to trace.",
    25: "The team studied issue-tracking and SAP Fiori patterns but kept IDTS within its approved scope. It is not a replacement for source control, CI/CD, sprint planning, or a full Jira platform.",
    27: "IDTS provides a focused learning and QA environment where role-aware CAP services and Fiori screens keep bug data, ownership, evidence, and history consistent. The immediate value is traceable collaboration rather than autonomous decision-making.",
    29: "Vision: provide a clear SAP-style defect workflow in which humans remain responsible for classification, assignment, lifecycle decisions, and acceptance, while the system validates data, records evidence, and offers bounded advisory assistance.",
    31: "In scope: authentication, bug creation, catalog validation, assignment/reassignment, lifecycle actions, comments, attachments, notifications, history, monitoring, and review-only AI suggestions. Out of scope: automatic bug fixing, autonomous AI workflow decisions, source-code management, production SAP transport management, and mentor/user sign-off performed by the team.",
    35: "Project management follows Jira work items, protected pull requests, role ownership, evidence-based QA gates, and repository/Drive document synchronization. Shared QA runs on Render with PostgreSQL, AWS S3 attachments, and Brevo transactional delivery.",
    38: "The implemented MVP covers AuthService, BugService, Fiori List Report/Object Page, dashboard/monitoring, collaboration, attachment storage, notification outbox, exact-action history, and four advisory AI review flows. Complexity is controlled by CAP/Fiori-supported patterns and a minimal provider seam.",
    40: "Objectives: preserve backend authority, prevent invalid or unauthorized writes, provide readable SAP Fiori UX, keep evidence auditable, and maintain truthful mentor documentation. Quality is measured through compile/build gates, programmatic regression, browser checks, persistence checks, and template validation.",
    42: "Key open risks are human UAT/sign-off, the decision for long-term PostgreSQL hosting, dependency vulnerabilities under IDTS-46, and disabled live OpenAI acceptance. These are disclosed and are not represented as completed work.",
    44: "The team uses short Jira-scoped branches, pull-request review, QA Depth Gate evidence, and staged documentation updates. Runtime changes and documentation remediation remain separate when their risk or approval path differs.",
    46: "Lifecycle: requirement and design review → implementation on a task branch → focused verification → PR review/gates → merge to dev → Shared QA deploy → acceptance evidence → SAP490 synchronization. Runtime remains the authority when a document is stale.",
    48: "Quality controls include backend authorization/validation tests, role-matrix browser checks, persistence/reload evidence, error-leak scans, secret scans, CAP/UI5 build gates, OfficeCLI validation, EN/VI parity, template fidelity, and Drive readback.",
    50: "Training uses ownership-specific beginner notes, Debug Labs, Knowledge Gates, and teach-back. The Knowledge Gate remains IN PROGRESS in a dedicated learning thread and is not claimed as passed in this report.",
    52: "Deliverables include source code, database models/seeds, deployment configuration, BRD/SRS/FRS, Blueprint, functional/technical specifications, configuration and change notes, diagrams, test workbooks, evidence, and this mentor-review report draft.",
    54: "Responsibility follows the team ownership above, with DonHV coordinating cross-layer integration and mentor-pack truth. Every task still requires review and evidence; ownership does not remove shared accountability.",
    56: "Team communication uses Jira for work tracking, GitHub for code review, repository PM/status files for handover, and Google Drive for mentor collaboration copies. Secrets and private credentials are excluded from these channels.",
    59: "Canonical Markdown and generated Office artifacts are versioned in Git. Google Drive files are updated in place to preserve IDs, parents, MIME types, and review links. Approval/signature fields are not completed by agents.",
    61: "Source changes use task branches and protected PRs into dev. Runtime deployments use a verified dev commit. Destructive Git commands and branch-protection bypass are prohibited for this workstream.",
    63: "Tools include SAP CAP/CDS, SAPUI5/Fiori Elements, Node.js, PostgreSQL, SQLite for local development, Render, AWS S3, Brevo, GitHub, Jira, OfficeCLI, LibreOffice, and controlled browser/API test harnesses.",
    66: "The SRS baseline is maintained separately in the official BRD/SRS/FRS package. This section summarizes the current implemented contract and links it to design and test evidence.",
    69: "Actors are Tester, Developer, PM, external email delivery, S3 object storage, and the optional AI provider seam. Fiori calls authenticated OData V4 services; CAP validates and persists data before external side effects are processed.",
    72: "Tester creates and follows defects; Developer reviews and resolves assigned work; PM monitors and reassigns work. Server-side authorization protects the same boundaries when callers bypass the UI.",
    75: "Functions include login/profile/logout; draft create and validation; assignment; lifecycle actions; comments; attachments; history; notification/email outbox; dashboard/workload; and Similar Bugs, Classification, Handoff, and Smart Assignment advisory reviews.",
    78: "Detailed functional specifications are provided in the Functional Specification workbook and FRS. Important workflows use exact OData actions and exact history action types rather than ambiguous generic labels.",
    79: "2.1 Defect lifecycle and collaboration",
    80: "2.1.1 Create, assign, process, retest, close, and reopen Bug",
    81: "Fiori draft creation uses NEW, repeated PATCH, and SAVE/CREATE. Editing an active Bug uses EDIT, PATCH, and UPDATE. Bound actions enforce role/status rules and produce history, next-processor, and notification side effects in the same controlled flow.",
    84: "The application keeps standard Fiori Elements List Report/Object Page behavior and uses SAPUI5 extensions only where richer interaction is required: login/profile shell, dashboard, smart assignment, collaboration, attachments, history pagination, and AI review dialogs.",
    85: "3.1 Defect management experience",
    86: "3.1.1 Bug List Report and Object Page",
    87: "a. Role-aware defect workspace",
    88: "Screens use role-aware actions, safe messages, readable labels, responsive layout, and no developer-facing implementation text. Backend checks remain authoritative even when a UI action is hidden.",
    91: "Related flows: authentication, Bug create/edit, assignment, lifecycle actions, collaboration, evidence, dashboard monitoring, notification, and advisory AI.",
    94: "Current UI evidence is stored under docs/pm/evidence and linked from the SAP490 test pack. Screenshots exclude credentials, tokens, private endpoints, and full recipient lists.",
    95: "See the Shared QA screenshots for AI review dialogs, assignment explanation, draft attachments, saved comments, and safe upload failure/recovery.",
    98: "UI fields and actions are derived from CDS annotations, manifest routing, OData metadata, and supported SAPUI5 extension APIs. Fixed code lists and backend-targeted errors prevent free-text corruption.",
    99: "See FRS and Functional Specification for field-level behavior and acceptance criteria.",
    103: "External interfaces: /odata/v4/auth/, /odata/v4/bug/, Render PostgreSQL, private AWS S3 object storage, Brevo transactional API/SMTP seam, and an optional disabled AI provider seam. Credentials are private environment configuration.",
    105: "Quality attributes include security, traceability, consistency, persistence, recoverability, readable SAP Fiori UX, portability between local SQLite and cloud PostgreSQL, and bounded external-provider failure.",
    107: "Appendices include the Blueprint, diagrams, specifications, test catalog, evidence index, risk/decision log, and Drive Mentor Index.",
    109: "Business rules cover role permissions, status transitions, assignee responsibility, next processor, mandatory reason/note fields, draft behavior, attachment constraints, notification isolation, and review-only AI.",
    111: "Common requirements: safe errors, no secrets, backend authority, history traceability, EN/VI document parity, evidence-based completion, and preserved Drive file identity.",
    113: "Application messages are cataloged in Functional and Technical Specification workbooks. UI messages must be safe and understandable; raw SQL, stack traces, credentials, and provider diagnostics are prohibited.",
    116: "The design baseline is described by the Blueprint and canonical diagram pack. This section provides a concise implementation-oriented summary.",
    119: "Architecture: SAPUI5/Fiori Elements client → OData V4 AuthService/BugService → CAP Node.js handlers → PostgreSQL. Attachments store metadata in PostgreSQL and binary content in S3. Email uses an outbox worker; advisory AI is isolated behind a provider seam and is currently disabled live.",
    121: "Repository packages are app/ for UI, srv/ for OData services and handlers, db/ for the CDS domain/seed data, scripts/qa for evidence harnesses, and docs/ for canonical knowledge, PM, BA, and SAP490 artifacts. This is a CAP package structure, not a classic ABAP package/TR.",
    123: "The CDS model contains Bugs, users/roles, classification catalogs, developer profiles/responsibilities, comments, attachments, history events/logs, notifications/deliveries, auth sessions, and AiSuggestions. Shared QA uses PostgreSQL; S3 stores attachment bytes.",
    125: "3.1 Defect lifecycle implementation",
    126: "Detailed designs use small CAP handlers for authorization, validation, transitions, history, notification, attachment orchestration, and AI normalization. Caller/callee and breakpoint guidance is maintained in docs/knowledge.",
    128: "Class/module relationships are represented by the component, service, and data-model diagrams in the Diagram Pack.",
    130: "Representative sequence: UI action → OData request → service.cds contract → service.js handler → validator/permission helper → CAP transaction → PostgreSQL/S3/outbox side effect → OData response → Fiori refresh.",
    131: "Authentication sequence: login UI → AuthService.login → credential verification → AuthSessions token → authenticated BugService requests.",
    132: "Other sequences are documented for lifecycle, attachments, notification, and advisory AI.",
    134: "3.2 Collaboration, notification, and advisory assistance",
    135: "Feature-level details are maintained in the FRS, Technical Specification, source knowledge mirrors, and canonical diagrams.",
    138: "Testing covers local-fast programmatic checks and Shared QA browser/API/persistence evidence. Human UAT and mentor approval remain pending.",
    140: "Scope includes authentication, validation, role boundaries, lifecycle, assignment, comments, attachments, history, dashboard, notifications, PostgreSQL/S3 persistence, and advisory AI. Live OpenAI provider acceptance is excluded because the provider is disabled by decision.",
    141: "Levels: unit/programmatic regression, integration/API, browser/system, and prepared UAT. Acceptance requires positive, negative/boundary, role, persistence, failure/recovery, and UI/UX evidence where applicable.",
    142: "Constraints: Shared QA contains demo data; active Bugs are not directly deletable through OData; six UAT cases require named human execution/sign-off; provider secrets are unavailable to public evidence.",
    144: "Strategy combines deterministic regression with falsification-first exploratory testing. Expected 400/401/403 negative responses are PASS when they prove validation or authorization.",
    146: "Types include unit, API/contract, integration, browser, persistence/reload, security/error-leak, email/outbox, object-storage, and document/template validation.",
    148: "Unit and integration checks run locally; system/browser and persistence checks run against frozen Shared QA; UAT remains a human acceptance activity.",
    150: "Tools: npm/CDS/UI5 scripts, controlled Node.js QA harnesses, browser automation, Render/Brevo evidence, OfficeCLI, LibreOffice, and document validators.",
    153: "DonHV consolidates evidence; NhanT owns QA review; SangVN and DatDT provide role-specific sign-off. In this run the automated multi-role rehearsal passed, while named human UAT signatures remain pending.",
    155: "Local environment uses Node.js/CAP with SQLite. Shared QA uses the frozen dev commit on Render, PostgreSQL, private S3, and Brevo. AI live provider is disabled.",
    157: "Milestones: local-fast gates → Shared QA lifecycle/AI/attachment/email smoke → document regeneration → Office/template/visual verification → Drive update → mentor review → human UAT/sign-off.",
    159: "The canonical catalog contains 27 planned cases. Current evidence supports 21 PASSED cases and 6 PREPARED UAT cases; none are represented as failed or blocked.",
    163: "Fresh results: 21/27 cases PASSED, including nine Shared QA functional scenarios; 6/27 UAT cases remain PREPARED. Pass rate among executed cases is 100%, but full-plan completion is not claimed. OpenAI live-provider acceptance is NOT ACCEPTED because it is disabled.",
    165: "This section lists the mentor-review release package and safe setup guidance. It is not a production release approval.",
    167: "Package: source repository, CDS model/services, Fiori application, Render configuration, QA scripts/evidence, BA documents, Blueprint, specifications, diagrams, test workbooks, change/configuration notes, PM matrices, and this report draft.",
    171: "Requirements: supported Node.js/npm, CAP dependencies, local SQLite or approved PostgreSQL, browser access, and private provider configuration for S3/Brevo. Do not commit secrets.",
    173: "Local: npm ci, deploy/refresh the local SQLite schema as documented, set approved local auth passwords privately, then npm start or cds watch. Shared QA is deployed from dev; database deploy/reset is not part of ordinary smoke testing.",
    177: "Users sign in, open the dashboard or Bug list, create/classify a Bug, optionally select an assignee, collaborate through comments/evidence, execute authorized lifecycle actions, and review history/notifications.",
    179: "3.2 Create and process a Bug",
    180: "Sign in as Tester/PM → Create Bug → enter valid classification/reproduction data → optionally attach safe evidence → Save → assign or leave Pending Assignment → Developer reviews/progresses/resolves → Tester/PM retests, closes, or reopens. History and notifications document each authorized step.",
    181: "3.3 Review advisory AI and monitoring",
    182: "Open Similar Bugs, Classification Suggestions, Smart Assignment Explanation, or Handoff Summary from its business section. Review the sanitized output; it never changes workflow automatically. PM dashboard and workload views support monitoring and navigation to the source Bugs.",
}


def set_cell(table, row, col, text):
    table.cell(row, col).text = text


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)

    # Preserve the official cover/table structure while replacing sample values.
    doc.tables[1].cell(2, 0).text = "Issue and Defect Tracking System in SAP (IDTS-SAP01)\nMentor Review Draft v0.1 — 2026-07-24"
    set_cell(doc.tables[2], 0, 0, "GSU26SAP01")
    set_cell(doc.tables[2], 0, 1, "GSU26SAP01")
    set_cell(doc.tables[2], 1, 0, "Group Members")
    set_cell(doc.tables[2], 1, 1, "DonHV; DatDT; SangVN; NhanT")
    set_cell(doc.tables[2], 2, 0, "Supervisor")
    set_cell(doc.tables[2], 2, 1, "Pending mentor confirmation")
    set_cell(doc.tables[2], 3, 0, "Ext Supervisor")
    set_cell(doc.tables[2], 3, 1, "N/A / pending confirmation")

    acronyms = [
        ("IDTS", "Issue and Defect Tracking System"), ("CAP", "SAP Cloud Application Programming Model"),
        ("CDS", "Core Data Services"), ("OData", "Open Data Protocol"), ("Fiori", "SAP user-experience design system"),
        ("UI5", "SAPUI5 web UI framework"), ("UAT", "User Acceptance Testing"), ("S3", "Amazon Simple Storage Service"),
        ("SMTP", "Simple Mail Transfer Protocol"), ("AI", "Artificial Intelligence; advisory-only in IDTS"),
        ("QA", "Quality Assurance"), ("PM", "Project Manager role"), ("TR", "SAP Transport Request; not used by this CAP project"),
    ]
    table = doc.tables[3]
    while len(table.rows) < len(acronyms) + 1:
        table.add_row()
    for row, (code, meaning) in enumerate(acronyms, 1):
        set_cell(table, row, 0, code)
        set_cell(table, row, 1, meaning)

    deliverables = [
        ("1", "Schedule/Task Tracking", "Jira Sprint 4 and PM handover/status files"),
        ("2", "Project Backlog", "Jira Epics/tasks with due dates, links, acceptance criteria, and evidence"),
        ("3", "Source Codes", "GitHub repository: CAP services, Fiori/UI5 app, QA and documentation generators"),
        ("4", "Database Script(s)", "CDS schema, seed datasets, PostgreSQL/SQLite deployment helpers"),
        ("5", "Final Report Document", "This official-template mentor-review draft; approval/signature pending"),
        ("6", "Requirements and Design", "BRD, SRS, FRS, Blueprint, Functional/Technical Specifications, diagrams"),
        ("7", "Testing", "27-case catalog, 12 test workbooks, Shared QA evidence, 21 PASSED and 6 PREPARED UAT"),
        ("8", "Deployment", "Render Shared QA, PostgreSQL, AWS S3, Brevo; private values excluded"),
        ("9", "User Guide / Knowledge", "Beginner-first knowledge mirrors, Debug Labs, and workflow guidance"),
    ]
    table = doc.tables[4]
    while len(table.rows) < len(deliverables) + 1:
        table.add_row()
    for row, values in enumerate(deliverables, 1):
        for col, value in enumerate(values):
            set_cell(table, row, col, value)

    for index, text in CONTENT.items():
        doc.paragraphs[index].text = text

    screenshots = [
        (94, ROOT / "docs/pm/evidence/idts-100/shared-qa-ai-browser/all-review-actions/01_similar_bugs_dialog.png"),
        (128, ROOT / "docs/pm/evidence/idts-100/shared-qa-ai-browser/all-review-actions/03_handoff_dialog.png"),
        (159, ROOT / "docs/pm/evidence/idts-100/shared-qa-attachments/idts73_saved_attachments_and_comments.png"),
    ]
    for paragraph_index, image in screenshots:
        if image.exists():
            doc.paragraphs[paragraph_index].add_run().add_picture(str(image), width=Inches(6.2))

    doc.core_properties.title = "IDTS-SAP01 Final Project Report — Mentor Review Draft"
    doc.core_properties.subject = "Evidence-backed mentor review; human UAT and approval pending"
    doc.core_properties.author = "IDTS SAP01 Team"
    doc.save(OUTPUT)
    print(OUTPUT.relative_to(ROOT))


if __name__ == "__main__":
    main()
